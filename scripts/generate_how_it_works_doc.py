#!/usr/bin/env python3
"""Word-документ: как работают ключевые системы (логика → код)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "TheMazeRunner_Как_это_работает.docx"


def setup(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def p(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def logic(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    r = para.add_run("Логика: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0, 70, 140)
    para.add_run(text)
    para.paragraph_format.space_after = Pt(8)


def code(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    r = para.add_run("Код: ")
    r.bold = True
    r.font.color.rgb = RGBColor(120, 0, 0)
    para.paragraph_format.space_after = Pt(4)
    block = doc.add_paragraph()
    run = block.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    block.paragraph_format.left_indent = Cm(0.5)
    block.paragraph_format.space_after = Pt(10)


def file_ref(doc: Document, path: str) -> None:
    para = doc.add_paragraph()
    r = para.add_run("Файл: ")
    r.bold = True
    para.add_run(path)
    para.paragraph_format.space_after = Pt(6)


def main() -> None:
    doc = Document()
    setup(doc)

    t = doc.add_heading("The Maze Runner — как это работает", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(doc,
      "Этот документ объясняет пять главных систем игры. "
      "Сначала — идея и логика простыми словами, потом — как это записано в коде и в каких файлах.")
    p(doc, f"Проект: {ROOT}")

    # ===== MONSTERS =====
    h1(doc, "1. NPC (розовые монстры) — как сделаны и как работают")

    h2(doc, "1.1. Как мы их «сделали» (от картинки до игры)")
    logic(doc,
          "Исходник — большой PNG monster_spritesheet.png (розовый персонаж на прозрачном фоне). "
          "На нём нарисованы три ряда кадров подряд: 6 кадров бега, 5 кадров «злого» вида с открывающимся ртом, "
          "5 кадров поимки игрока. Python-скрипт extract_monster_pink_sheet.py находит границы кадров по прозрачности "
          "и собирает аккуратный лист monster_pink_sheet.png размером 6 колонок × 3 ряда. "
          "В игре класс MonsterPinkSheet режет этот лист на прямоугольники, а класс Monster выбирает нужный кадр.")
    file_ref(doc, "scripts/extract_monster_pink_sheet.py → assets/monster_pink_sheet.png")
    file_ref(doc, "src/MonsterPinkSheet.hpp, src/MonsterPinkSheet.cpp")
    file_ref(doc, "src/Monster.hpp, src/Monster.cpp")

    h2(doc, "1.2. Два типа поведения (ИИ)")
    logic(doc,
          "Patrol (патруль) — монстр ходит по коридору в заданном направлении (влево/вправо/вверх/вниз). "
          "Уперся в стену — разворачивается. Если включён lockPatrolAxis, он только разворачивается на 180° "
          "и не «сворачивает» на перекрёстках — это нужно, чтобы монстр оставался на одной линии коридора.")
    logic(doc,
          "Chase (погоня) — пока игрок далеко, монстр ведёт себя как патруль. "
          "Когда игрок подходит ближе CHASE_START_DISTANCE (220 px), монстр включает погоню: "
          "каждый кадр спрашивает у лабиринта «куда идти дальше по кратчайшему пути» (алгоритм BFS) "
          "и делает один шаг к игроку. Если игрок отбежал дальше CHASE_STOP_DISTANCE (140 px), погоня прекращается, "
          "монстр успокаивается и снова патрулирует.")
    file_ref(doc, "src/LevelData.hpp — MonsterAI::Patrol / Chase, поля patrolDirX/Y, lockPatrolAxis")
    file_ref(doc, "src/Constants.hpp — CHASE_START_DISTANCE, CHASE_STOP_DISTANCE, MONSTER_CHASE_FACTOR")
    code(doc,
         "// LevelData.hpp — пример спавна chase-монстра на ур. 3:\n"
         "M{ 7, 13, 1, 0, MonsterAI::Chase, 1.05f, true }\n"
         "// gridX=7, gridY=13, патруль вправо, chase, скорость ×1.05, lockPatrolAxis")

    h2(doc, "1.3. Анимации монстра")
    logic(doc,
          "Run — обычный розовый бег (6 кадров, циклически). "
          "Transform — когда chase-монстр впервые замечает игрока, проигрывается 5 кадров превращения в злого. "
          "Evil — злой вид: тот же ряд спрайтов, но выбирается колонка по дистанции до игрока "
          "(чем ближе — тем шире открыт рот, колонки 1–4). "
          "Kill — при столкновении с игроком: 4 кадра «съел», потом Game Over.")
    code(doc,
         "void Monster::onChaseStarted() {\n"
         "    animState_ = MonsterAnimState::Transform;  // начало злости\n"
         "}\n"
         "int Monster::mouthColFromDist(float dist) const {\n"
         "    if (dist > 220.f) return 0;\n"
         "    if (dist > 170.f) return 1;  // рот всё шире\n"
         "    ...\n"
         "    return 4;\n"
         "}")

    h2(doc, "1.4. Движение и коллизии")
    logic(doc,
          "Монстр — круг с радиусом 12 px. Движение такое же, как у игрока: maze.moveCircle() "
          "сначала пробует сдвинуть по X, потом по Y — получается скольжение вдоль стен в углах. "
          "Для chase используется nextStepBfs: из текущей клетки сетки BFS ищет путь до клетки игрока "
          "и возвращает соседнюю клетку — первый шаг пути.")
    code(doc,
         "// Monster.cpp — chase:\n"
         "const sf::Vector2i nextG = maze.nextStepBfs(fromG, toG);\n"
         "const sf::Vector2f target = maze.gridCenter(nextG);\n"
         "maze.moveCircle(pos_, dir * speed_ * dt, radius_);\n\n"
         "// Game.cpp — столкновение с игроком:\n"
         "if (circlesOverlap(player_.position(), player_.radius(),\n"
         "                   monster.position(), monster.radius())) {\n"
         "    monsters_[i].startKill();\n"
         "    state_ = GameState::Caught;\n"
         "}")

    h2(doc, "1.5. Где задаются монстры на уровнях")
    logic(doc,
          "Все уровни описаны в buildLevels() в LevelData.hpp. "
          "Каждый монстр — структура MonsterSpawn: координата на сетке, направление патруля, тип AI, множитель скорости.")
    file_ref(doc, "src/LevelData.hpp — функция buildLevels()")
    file_ref(doc, "src/Game.cpp — startLevel(): цикл for (spawn : lvl.monsters) → m.spawn()")

    # ===== MAP TEXTURES =====
    h1(doc, "2. Карта лабиринта с текстурами")

    h2(doc, "2.1. Логика: от символов к картинке на экране")
    logic(doc,
          "Лабиринт — это массив строк в LevelData.hpp. Символ # — стена, . — пол, S — старт, E — выход. "
          "При загрузке Maze::loadFromLayout() запоминает координаты S и E в пикселях и заменяет их на точку. "
          "Каждый кадр Maze::draw() проходит по всем клеткам: для стены рисует кусок wallN.jpg, "
          "для пола — кусок floor_tile.jpg. Если JPG не загрузился — рисуются цветные квадраты-заглушки.")
    logic(doc,
          "Размер клетки (tileSize) у каждого уровня свой (24–32 px) или подгоняется под экран функцией "
          "fitTileSizeForLevel(), чтобы большой лабиринт влез на монитор. "
          "applyGameView() масштабирует камеру (sf::View), чтобы весь лабиринт + полоска HUD были видны.")

    h2(doc, "2.2. Как «нарезается» текстура на клетки")
    logic(doc,
          "Одна JPG большая (пол или стена). Чтобы не растягивать всю картинку на одну клетку, "
          "функция drawTiledTexture() для клетки (x, y) вырезает из JPG прямоугольник, смещённый "
          "в зависимости от координат сетки. FLOOR_TEXTURE_ZOOM=16 и WALL_TEXTURE_ZOOM=8 задают, "
          "какой кусок исходника брать — чем больше zoom, тем мельче «плитка» и крупнее узор на полу.")
    file_ref(doc, "src/Maze.cpp — drawTiledTexture(), Maze::draw()")
    file_ref(doc, "src/Constants.hpp — FLOOR_TEXTURE_ZOOM, WALL_TEXTURE_ZOOM")
    code(doc,
         "void drawTiledTexture(..., int gridX, int gridY, ..., float zoom) {\n"
         "    const int srcCellW = texW / zoom;\n"
         "    const int srcX = (gridX * srcCellW) % texW;  // сдвиг по сетке\n"
         "    sprite.setTextureRect(sf::IntRect(srcX, srcY, cropW, cropH));\n"
         "    sprite.setScale(ts / cropW, ts / cropH);     // на размер клетки\n"
         "}")

    h2(doc, "2.3. Откуда взялись картинки и как подключили")
    logic(doc,
          "floor_tile.jpg — одна текстура пола для всех уровней. "
          "wall1.jpg … wall5.jpg — своя стена для каждого из 5 уровней. "
          "Файлы лежат в assets/. Art.cpp загружает их в ArtPack. "
          "В Game::startLevel() для текущего уровня вызывается maze_.setWallTexture(&art_.wallTiles[currentLevel_]).")
    file_ref(doc, "src/Art.cpp — загрузка floor_tile.jpg и wall1–5.jpg")
    file_ref(doc, "src/Game.cpp — startLevel(): setFloorTexture, setWallTexture")
    code(doc,
         "// Game.cpp\n"
         "maze_.setFloorTexture(art_.floorTileLoaded ? &art_.floorTile : nullptr);\n"
         "maze_.setWallTexture(&art_.wallTiles[currentLevel_]);")

    h2(doc, "2.4. Выход (портал) на клетке E")
    logic(doc,
          "Отдельно от пола/стен: на позиции exitPos_ рисуется ExitPortal — "
          "анимация из portal_sheet.png (64 кадра из GIF). Это не часть сетки Maze::draw().")
    file_ref(doc, "src/ExitPortal.cpp, src/PortalSheet.cpp")

    # ===== MUSIC =====
    h1(doc, "3. Музыка и звуки")

    h2(doc, "3.1. Логика")
    logic(doc,
          "SFML Audio умеет проигрывать MP3 через sf::Music (потоковое, для длинных треков) "
          "и sf::Sound + sf::SoundBuffer (короткие эффекты). "
          "В меню играет menu_music.mp3 по кругу. На уровне — maze_music.mp3 по кругу: "
          "если уровень длиннее трека, музыка начинается сначала (setLoop(true)). "
          "Если уровень короче — музыка просто играет до конца уровня, при выходе stopLevelMusic(). "
          "При победе — victory_sound.mp3, при поражении — defeat_sound.mp3; фоновая музыка уровня останавливается.")
    file_ref(doc, "assets/menu_music.mp3, maze_music.mp3, victory_sound.mp3, defeat_sound.mp3")

    h2(doc, "3.2. Код загрузки и переключения")
    code(doc,
         "// Game.cpp — загрузка (пробует assets/ и ../assets/):\n"
         "levelMusic_.openFromFile(\"assets/maze_music.mp3\");\n"
         "levelMusic_.setLoop(true);\n"
         "levelMusic_.play();\n\n"
         "// При старте уровня:\n"
         "void Game::startLevelMusic() {\n"
         "    stopMenuMusic();\n"
         "    levelMusic_.setLoop(true);\n"
         "    levelMusic_.play();\n"
         "}\n\n"
         "// При поражении:\n"
         "void Game::playDefeatSound() {\n"
         "    stopLevelMusic();\n"
         "    defeatSound_.play();\n"
         "}")

    h2(doc, "3.3. Подключение SFML Audio в проекте")
    logic(doc,
          "В TheMazeRunner.vcxproj в AdditionalDependencies добавлен sfml-audio. "
          "PostBuild копирует sfml-audio-2.dll и openal32.dll рядом с exe — без них звук не работает.")
    file_ref(doc, "TheMazeRunner.vcxproj — Link + PostBuildEvent")

    # ===== MAIN MENU =====
    h1(doc, "4. Главное меню")

    h2(doc, "4.1. Логика")
    logic(doc,
          "При запуске state_ = MainMenu. В render() вызывается drawMainMenu() вместо лабиринта. "
          "Слои отрисовки снизу вверх: 1) фон menu_bg.png на весь экран (cover, как обои); "
          "2) анимация погони MenuChase; 3) список уровней неоновым текстом; 4) подсказка по клавишам внизу. "
          "Выбор уровня — стрелки ↑↓ или цифры 1–5, старт — Enter. "
          "Название «THE MAZE RUNNER» нарисовано на самом фоне (menu_bg.png), в коде отдельно не рисуется.")
    logic(doc,
          "Текст меню фиолетовый неон (#D050FF): обводка + лёгкое свечение у выбранного пункта. "
          "Шрифт — Arial из Windows (C:/Windows/Fonts/arial.ttf). Русский текст через utf8().")

    h2(doc, "4.2. Код меню")
    file_ref(doc, "src/Game.cpp — drawMainMenu(), drawMenuBackground(), drawNeonMenuText(), processEvents()")
    code(doc,
         "void Game::drawMainMenu() {\n"
         "    drawMenuBackground();                    // menu_bg.png\n"
         "    menuChase_.draw(window_, winH * 0.38f, winW);  // погоня\n"
         "    // список уровней на y = winH * 0.54f\n"
         "    for (уровни) drawNeonMenuText(..., selected);\n"
         "}\n\n"
         "void Game::drawNeonMenuText(sf::Text& text, bool selected) {\n"
         "    text.setFillColor(sf::Color(208, 80, 255));  // #D050FF\n"
         "    text.setOutlineThickness(selected ? 4.f : 2.5f);\n"
         "    // glow-копия текста для выбранного пункта\n"
         "}")

    h2(doc, "4.3. События клавиатуры в меню")
    code(doc,
         "if (event.key.code == sf::Keyboard::Up)\n"
         "    menuSelectedLevel_ = (menuSelectedLevel_ - 1 + n) % n;\n"
         "if (event.key.code == sf::Keyboard::Enter)\n"
         "    startLevel(menuSelectedLevel_);  // → Playing, музыка уровня")

    # ===== MENU CHASE =====
    h1(doc, "5. Погоня на главном меню (монстр за игроком)")

    h2(doc, "5.1. Логика — что видит игрок")
    logic(doc,
          "Между названием на фоне и списком уровней бежит игрок слева направо, "
          "а за ним на расстоянии 500 px — злой розовый монстр. "
          "Оба используют те же спрайт-листы, что и в игре (player_sheet, monster_pink_sheet), "
          "но только кадры «бег вправо» у игрока и «злой рот открыт» у монстра. "
          "Когда игрок убегает за правый край экрана, он исчезает и ждёт: монстр добегает до края сам, "
          "и только тогда оба появляются слева снова — так не бывает, что игрок уже слева, а монстр ещё справа.")
    logic(doc,
          "Высота спрайтов ~8.8% ширины экрана. Ноги игрока и монстра выровнены по нижнему краю "
          "(разные точки origin Y: 0.55 у игрока, 0.70 у монстра + формула alignedMonsterY).")

    h2(doc, "5.2. Логика — цикл движения")
    bullet(doc, "Обычный режим: playerX_ и monsterX_ растут со скоростью 240 px/с")
    bullet(doc, "Монстр всегда отстаёт на MONSTER_OFFSET = 500 px")
    bullet(doc, "Игрок вышел за правый край → waitingForMonster_ = true, игрок скрыт")
    bullet(doc, "Монстр добежал за край → оба телепортируются: игрок -70, монстр -570")
    bullet(doc, "Анимация: игрок переключает RUN_A/RUN_B каждые 0.11 с; монстр — кадры злого ряда")

    h2(doc, "5.3. Код MenuChase")
    file_ref(doc, "src/MenuChase.hpp, src/MenuChase.cpp")
    file_ref(doc, "src/Game.cpp — init(): menuChase_.setSheets(); update(): menuChase_.update()")
    code(doc,
         "void MenuChase::update(float dt, float laneWidth) {\n"
         "    if (waitingForMonster_) {\n"
         "        monsterX_ += SPEED * dt;\n"
         "        if (monsterX_ > right) { /* сброс слева */ }\n"
         "    } else {\n"
         "        playerX_ += SPEED * dt;\n"
         "        monsterX_ += SPEED * dt;\n"
         "        if (playerX_ - monsterX_ < MONSTER_OFFSET)\n"
         "            monsterX_ = playerX_ - MONSTER_OFFSET;\n"
         "        if (playerX_ > right) waitingForMonster_ = true;\n"
         "    }\n"
         "}")

    h2(doc, "5.4. Код отрисовки спрайтов в погоне")
    code(doc,
         "// Игрок — колонка COL_RIGHT, ряды RUN_A / RUN_B:\n"
         "sprite.setTextureRect(playerSheet_->cellRect(COL_RIGHT, row));\n\n"
         "// Монстр — ряд ROW_TRANSFORM, последняя колонка (рот открыт):\n"
         "int col = TRANSFORM_FRAMES - 1;\n"
         "sprite.setTextureRect(monsterSheet_->cellRect(col, ROW_TRANSFORM));\n\n"
         "// Выравнивание ног:\n"
         "float alignedMonsterY(float centerY, float displayH) {\n"
         "    float playerBottom = centerY + (1 - 0.55) * displayH;\n"
         "    return playerBottom - (1 - 0.70) * monsterH;\n"
         "}")

    h2(doc, "5.5. Константы MenuChase")
    bullet(doc, "SPEED = 240, MONSTER_OFFSET = 500, EDGE_PAD = 70")
    bullet(doc, "Позиция Y в меню: winH * 0.38f (Game::drawMainMenu)")
    bullet(doc, "Список уровней: winH * 0.54f")

    # ===== SUMMARY =====
    h1(doc, "6. Как всё связано в одном кадре меню")
    p(doc, "Game::run() → update(): если MainMenu, только menuChase_.update(dt).")
    p(doc, "Game::run() → render(): drawMainMenu() → фон → погоня → текст.")
    p(doc, "Enter → startLevel() → stopMenuMusic(), startLevelMusic(), state_ = Playing.")

    h1(doc, "7. Краткая шпаргалка по файлам")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Система"
    table.rows[0].cells[1].text = "Главные файлы"
    for a, b in [
        ("Монстры (NPC)", "Monster.cpp, LevelData.hpp, MonsterPinkSheet, extract_monster_pink_sheet.py"),
        ("Карта + текстуры", "Maze.cpp, LevelData.hpp, Art.cpp, Constants.hpp"),
        ("Музыка", "Game.cpp (load/start/stop Music и Sound)"),
        ("Главное меню", "Game.cpp (drawMainMenu, processEvents)"),
        ("Погоня в меню", "MenuChase.cpp, Game.cpp (init, drawMainMenu)"),
    ]:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b

    p(doc, "— Конец документа —")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
