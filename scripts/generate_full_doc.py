#!/usr/bin/env python3
"""Генерация полного описания проекта The Maze Runner в Word."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "TheMazeRunner_Полное_описание.docx"


def h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def p(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def code(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def main() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_heading("The Maze Runner — полное описание проекта", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(doc, "Курсовой проект: игра-лабиринт на C++17 и SFML 2.6.1 (Visual Studio 2022, x64).")
    p(doc, f"Папка проекта: {ROOT}")
    p(doc, "Документ описывает структуру файлов, логику игры и конкретный код.")

    # 1
    h1(doc, "1. Общая идея и архитектура")
    p(doc,
      "Игрок выбирает уровень в главном меню, бежит по лабиринту от старта (S) к выходу (E), "
      "уклоняясь от монстров и укладываясь в таймер. При победе — звук победы и переход дальше; "
      "при поимке монстром или истечении времени — звук поражения, анимация смерти, экран Game Over.")
    p(doc, "Главный цикл игры (класс Game):")
    bullet(doc, "processEvents() — клавиатура, меню, Esc, Enter")
    bullet(doc, "update(dt) — логика по состоянию GameState")
    bullet(doc, "render() — отрисовка кадра")
    p(doc, "Точка входа — src/main.cpp: создаёт Game, вызывает init() и run().")

    h2(doc, "Состояния игры (enum GameState в src/Game.hpp)")
    bullet(doc, "MainMenu — меню с фоном, списком уровней и анимацией погони")
    bullet(doc, "Playing — активный уровень")
    bullet(doc, "Caught — монстр поймал игрока, проигрывается kill-анимация монстра")
    bullet(doc, "Dying — анимация смерти игрока")
    bullet(doc, "LevelComplete — уровень пройден, ждём Enter")
    bullet(doc, "GameComplete — все 5 уровней пройдены")
    bullet(doc, "GameOver — поражение, ждём Enter/Esc/M")

    # 2
    h1(doc, "2. Структура папок проекта")
    bullet(doc, "src/ — весь исходный код C++")
    bullet(doc, "assets/ — картинки, музыка, звуки (копируются в папку exe при сборке)")
    bullet(doc, "scripts/ — Python-скрипты для нарезки спрайтов и проверки лабиринтов")
    bullet(doc, "docs/ — документация (этот файл, AGENT_HANDOFF.md, Kursovaya.rtf)")
    bullet(doc, "third_party/SFML/ — библиотека SFML (ставится через SETUP.bat)")
    bullet(doc, "x64/Debug/ — собранный exe и DLL после сборки")
    bullet(doc, "TheMazeRunner.sln / .vcxproj — решение Visual Studio")
    bullet(doc, "sfml_config.props — путь к SFML для MSBuild")
    bullet(doc, "SETUP.bat, ИГРАТЬ.bat — установка SFML и быстрый запуск")

    # 3
    h1(doc, "3. Файлы исходного кода (src/)")

    h2(doc, "3.1. main.cpp")
    p(doc, "Минимальная точка входа. Создаёт объект Game, инициализирует окно и ресурсы. "
          "Если init() вернул false — выход с кодом 1. Иначе game.run() крутит цикл, пока окно открыто.")

    h2(doc, "3.2. Constants.hpp")
    p(doc, "Глобальные константы баланса и отображения:")
    bullet(doc, "TILE_SIZE = 32 — базовый размер клетки")
    bullet(doc, "WINDOW_MARGIN = 80 — полоса под HUD снизу")
    bullet(doc, "SCREEN_WIDTH_RATIO / HEIGHT_RATIO — доля экрана под лабиринт")
    bullet(doc, "PLAYER_SPEED = 130 — скорость игрока (px/с)")
    bullet(doc, "MONSTER_CHASE_FACTOR = 0.82 — преследующий монстр чуть медленнее игрока")
    bullet(doc, "CHASE_START_DISTANCE = 220, CHASE_STOP_DISTANCE = 140 — гистерезис погони")
    bullet(doc, "COLLISION_RADIUS_SCALE = 0.88 — чуть меньший радиус для плавных углов")
    bullet(doc, "FLOOR_TEXTURE_ZOOM = 16, WALL_TEXTURE_ZOOM = 8 — масштаб текстур пола/стен")

    h2(doc, "3.3. Utf.hpp")
    p(doc, "Функция utf8(std::string) переводит UTF-8 строку в sf::String. "
          "Нужна для русского текста в меню и HUD. Важно: не вызывать utf8() дважды на одной строке.")

    h2(doc, "3.4. LevelData.hpp")
    p(doc, "Все данные уровней — лабиринты, таймеры, монстры.")
    bullet(doc, "MonsterAI::Patrol — ходит по коридору туда-обратно")
    bullet(doc, "MonsterAI::Chase — патрулирует, пока игрок далеко; при приближении преследует по BFS")
    bullet(doc, "MonsterSpawn — координаты сетки, направление патруля, множитель скорости, lockPatrolAxis")
    bullet(doc, "LevelConfig — имя, layout (строки символов), timeLimitSeconds, baseMonsterSpeed, tileSize")
    bullet(doc, "scaleLayout(layout, factor) — удваивает каждую клетку (уровень 1)")
    bullet(doc, "buildLevels() — возвращает vector из 5 уровней")
    p(doc, "Символы в layout: # стена, . пол, S старт, E выход. После загрузки S и E заменяются на точку.")

    h3(doc, "Уровни (кратко)")
    bullet(doc, "Ур.1 — разведка, 1 патрульный монстр, 90 с, tileSize 24")
    bullet(doc, "Ур.2 — 2 патруля (горизонталь + вертикаль с lockPatrolAxis), 75 с")
    bullet(doc, "Ур.3 — 2 патруля + 1 chase на (7,13), 60 с")
    bullet(doc, "Ур.4 — 3 chase-монстра, 55 с")
    bullet(doc, "Ур.5 — 3 chase с горизонтальным патрулём, 40 с, финал")

    h2(doc, "3.5. Maze.hpp / Maze.cpp")
    p(doc, "Лабиринт: сетка символов, коллизии, отрисовка пола/стен, BFS для ИИ монстров.")
    bullet(doc, "loadFromLayout() — парсит layout, находит S/E, выравнивает ширину строк")
    bullet(doc, "moveCircle(pos, delta, radius) — движение с раздельной проверкой X и Y (скольжение вдоль стен)")
    bullet(doc, "isWallAtPixel() — круг vs прямоугольники клеток-стен")
    bullet(doc, "nextStepBfs(from, to) — следующий шаг кратчайшего пути для chase-монстра")
    bullet(doc, "nearestFloorCenter() — если спавн в стене, ищет ближайший проход")
    bullet(doc, "draw() — для каждой клетки: стена (JPG или серый квадрат), пол (JPG или синеватый)")
    bullet(doc, "drawTiledTexture() — вырезает кусок большой JPG под координаты клетки (эффект «плитки»)")

    h2(doc, "3.6. Player.hpp / Player.cpp")
    p(doc, "Игрок: ввод WASD/стрелки, движение, анимация из player_sheet.png.")
    bullet(doc, "handleInput() — читает клавиши, нормализует вектор направления, сохраняет lastFacing_")
    bullet(doc, "update() — Idle или Run, двигает через maze.moveCircle(), advanceAnim()")
    bullet(doc, "directionToColumn() — 8 направлений → колонка спрайт-листа (COL_FRONT, COL_RIGHT, …)")
    bullet(doc, "advanceAnim() — Run: кадры A/B каждые 0.11 с; Dying: 4 кадра смерти с разными таймингами")
    bullet(doc, "drawSprite() — берёт cellRect(facingCol_, row) из PlayerSheet")
    bullet(doc, "Fallback: если лист не загрузился — синий процедурный круг из Art::player")

    h2(doc, "3.7. PlayerSheet.hpp / PlayerSheet.cpp")
    p(doc, "Спрайт-лист игрока: 8 колонок (направления) × 7 рядов (анимации).")
    bullet(doc, "Ряды: IDLE, RUN_A, RUN_B, FALL, DEAD_BACK, DEAD_DOWN_A, DEAD_DOWN_B")
    bullet(doc, "loadFromFile() — грузит PNG, делит на cellW × cellH")
    bullet(doc, "cellRect(col, row) — прямоугольник кадра в текстуре")

    h2(doc, "3.8. Monster.hpp / Monster.cpp")
    p(doc, "Розовые монстры с тремя режимами анимации на monster_pink_sheet.png.")
    bullet(doc, "spawn() — ставит в nearestFloorCenter, задаёт AI и скорость")
    bullet(doc, "updatePatrol() — движение по patrolDir; при стене разворот или смена направления")
    bullet(doc, "lockPatrolAxis — только разворот на 180°, не сворачивает на перекрёстках")
    bullet(doc, "updateChase() — дистанция до игрока; chasing_ с гистерезисом; BFS шаг к игроку")
    bullet(doc, "onChaseStarted() — запуск анимации Transform (превращение в злого)")
    bullet(doc, "onChaseCalmed() — сброс evil_ при отдалении игрока")
    bullet(doc, "mouthColFromDist() — колонка «рот открыт» по дистанции (0–4)")
    bullet(doc, "startKill() / Kill — 4 кадра поимки игрока")
    bullet(doc, "drawSprite() — зеркало по facingLeft_ через отрицательный scaleX")

    h2(doc, "3.9. MonsterPinkSheet.hpp / MonsterPinkSheet.cpp")
    bullet(doc, "6×3 сетка: ROW_RUN (6 кадров бега), ROW_TRANSFORM (5 кадров злого), ROW_KILL (5, в игре 4)")
    bullet(doc, "Нарезка из assets/monster_pink_sheet.png (скрипт extract_monster_pink_sheet.py)")

    h2(doc, "3.10. ExitPortal.hpp / ExitPortal.cpp")
    p(doc, "Анимированный выход на клетке E. Крутит кадры portal_sheet.png (64 кадра, 30 мс). "
          "Fallback — зелёный круг exitMarker из Art. Размер ≈ tileSize × 1.4.")

    h2(doc, "3.11. PortalSheet.hpp / PortalSheet.cpp")
    p(doc, "Лист 8×8 кадров портала из GIF XDZT (скрипт extract_portal_sheet.py). FRAME_TIME = 0.03 с.")

    h2(doc, "3.12. MenuChase.hpp / MenuChase.cpp")
    p(doc, "Декоративная погоня в главном меню: игрок бежит вправо, за ним злой монстр.")
    bullet(doc, "SPEED = 240 px/s, MONSTER_OFFSET = 500 px")
    bullet(doc, "Игрок: COL_RIGHT, ROW_RUN_A/B; монстр: ROW_TRANSFORM, последняя колонка")
    bullet(doc, "waitingForMonster_ — игрок уходит за правый край и ждёт, пока монстр тоже выйдет")
    bullet(doc, "alignedMonsterY() — выравнивание «ног» игрока и монстра")
    bullet(doc, "Позиция Y = winH × 0.38 (между названием на фоне и списком уровней)")

    h2(doc, "3.13. Art.hpp / Art.cpp")
    p(doc, "ArtPack — все загружаемые ресурсы. makeArt() вызывается один раз в Game::init().")
    bullet(doc, "playerSheet, monsterPinkSheet, portalSheet — спрайт-листы")
    bullet(doc, "floorTile, wallTiles[5], menuBackground — JPG/PNG текстуры")
    bullet(doc, "player — процедурный синий круг (fallback игрока)")
    bullet(doc, "exitMarker — процедурный зелёный круг (fallback портала)")
    bullet(doc, "makeSoftCircle() / makeExit() — генерация текстур в памяти без файлов")

    h2(doc, "3.14. Game.hpp / Game.cpp")
    p(doc, "Центральный класс — окно, состояние, звук, уровни, отрисовка.")
    h3(doc, "Инициализация init()")
    bullet(doc, "buildLevels(), шрифт Arial, музыка меню/уровня, звуки победы/поражения")
    bullet(doc, "makeArt(), menuChase_.setSheets()")
    bullet(doc, "Полноэкранное окно, VSync, лимит 120 FPS, startMenuMusic()")

    h3(doc, "startLevel(index)")
    bullet(doc, "fitTileSizeForLevel() — подгоняет tileSize под экран")
    bullet(doc, "maze_.loadFromLayout(), текстуры пола/стены по номеру уровня")
    bullet(doc, "Спавн игрока, монстров, портала; сброс таймера; Playing + музыка уровня")

    h3(doc, "update(dt) — логика по состоянию")
    bullet(doc, "MainMenu → menuChase_.update()")
    bullet(doc, "Caught → kill-анимация монстра → Dying")
    bullet(doc, "Dying → смерть игрока → GameOver")
    bullet(doc, "Playing → ввод, монстры, таймер, коллизия с монстрами, проверка выхода")

    h3(doc, "render()")
    bullet(doc, "MainMenu → фон menu_bg.png, MenuChase, неоновый список уровней (#D050FF)")
    bullet(doc, "Playing → applyGameView(), лабиринт, портал, монстры, игрок, HUD")
    bullet(doc, "Overlay — затемнение + текст для LevelComplete / GameOver")

    h3(doc, "applyGameView()")
    p(doc, "Масштабирует вид так, чтобы весь лабиринт + HUD влезли на экран с сохранением пропорций (letterbox).")

    h3(doc, "Звук")
    bullet(doc, "menu_music.mp3 — в меню, loop")
    bullet(doc, "maze_music.mp3 — на уровне, loop (длиннее уровня — начинается сначала)")
    bullet(doc, "defeat_sound.mp3 — при поражении, вместо музыки")
    bullet(doc, "victory_sound.mp3 — при достижении выхода")

  # 4
    h1(doc, "4. Ассеты (assets/)")
    p(doc, "При сборке PostBuild в TheMazeRunner.vcxproj копирует файлы в x64/Debug/assets/.")
    bullet(doc, "player_sheet.png — игрок (нарезанный лист)")
    bullet(doc, "player_spritesheet.png — исходник для scripts/extract_player_sheet.py (в игру не грузится)")
    bullet(doc, "monster_pink_sheet.png — розовые монстры")
    bullet(doc, "monster_spritesheet.png — исходник для extract_monster_pink_sheet.py")
    bullet(doc, "portal_sheet.png — анимация выхода")
    bullet(doc, "floor_tile.jpg — текстура пола")
    bullet(doc, "wall1.jpg … wall5.jpg — текстура стен по уровням")
    bullet(doc, "menu_bg.png — фон главного меню")
    bullet(doc, "maze_music.mp3, menu_music.mp3, defeat_sound.mp3, victory_sound.mp3")

    # 5
    h1(doc, "5. Скрипты (scripts/)")
    bullet(doc, "extract_player_sheet.py — нарезка player_sheet из player_spritesheet.png")
    bullet(doc, "extract_monster_pink_sheet.py — нарезка monster_pink_sheet")
    bullet(doc, "extract_portal_sheet.py — GIF → portal_sheet.png")
    bullet(doc, "check_maze.py, gen_square_maze.py, square_levels.py — генерация/проверка лабиринтов")
    bullet(doc, "setup_sfml.ps1 — используется SETUP.bat")
    bullet(doc, "extract_maze.py — вспомогательный")

    # 6
    h1(doc, "6. Сборка и запуск")
    bullet(doc, "SETUP.bat — распаковка SFML в third_party/SFML")
    bullet(doc, "Visual Studio: TheMazeRunner.sln, x64 Debug, Ctrl+F5")
    bullet(doc, "PostBuild копирует sfml-*.dll, openal32.dll и assets/")
    bullet(doc, "ИГРАТЬ.bat — запуск уже собранного exe")

    # 7
    h1(doc, "7. Поток данных (схема)")
    p(doc, "main → Game::init → makeArt + buildLevels")
    p(doc, "Game::run → каждый кадр: события → update → render")
    p(doc, "startLevel → Maze + Player + Monsters + ExitPortal + музыка")
    p(doc, "Playing: Player.handleInput → Maze.moveCircle; Monster.update → BFS/patrol")
    p(doc, "Победа: дистанция до exitPos < tileSize×0.45 → victory sound → LevelComplete")
    p(doc, "Поражение: circlesOverlap игрок-монстр → Caught → Kill → Dying → GameOver")

    # 8
    h1(doc, "8. Управление")
    bullet(doc, "Меню: ↑↓ выбор, 1–5 быстрый выбор, Enter старт, Esc выход/назад")
    bullet(doc, "Игра: WASD / стрелки — движение")
    bullet(doc, "Esc / M — в меню; Enter — следующий уровень после победы")

    # 9
    h1(doc, "9. Ключевые фрагменты кода")

    h2(doc, "Точка входа (main.cpp)")
    code(doc, "Game game;\nif (!game.init()) return 1;\ngame.run();")

    h2(doc, "Коллизия игрок–монстр (Game.cpp)")
    code(doc, "circlesOverlap(player_.position(), player_.radius(),\n"
              "              monster.position(), monster.radius())")

    h2(doc, "Движение со скольжением (Maze.cpp)")
    code(doc, "if (!isWallAtPixel(newX, pos.y, cr)) pos.x = newX;\n"
              "if (!isWallAtPixel(pos.x, newY, cr)) pos.y = newY;")

    h2(doc, "8 направлений игрока (Player.cpp)")
    code(doc, "if (std::abs(d.x) < 0.45f)\n"
              "    return d.y > 0.f ? COL_FRONT : COL_BACK;")

    h2(doc, "Chase-монстр: шаг BFS (Monster.cpp)")
    code(doc, "const sf::Vector2i nextG = maze.nextStepBfs(fromG, toG);\n"
              "maze.moveCircle(pos_, dir * speed_ * dt, radius_);")

    # 10
    h1(doc, "10. Типичные проблемы")
    bullet(doc, "LNK1104 — закройте exe перед пересборкой")
    bullet(doc, "OpenAL32.dll — копируется PostBuild из SFML/bin")
    bullet(doc, "Квадраты вместо русского — не двойной utf8()")
    bullet(doc, "Монстр сходит с линии патруля — включить lockPatrolAxis в LevelData")

    h1(doc, "11. Таблица: какой файл за что отвечает")
    rows = [
        ("src/main.cpp", "Точка входа программы"),
        ("src/Game.hpp / Game.cpp", "Окно, цикл игры, меню, звук, HUD, состояния"),
        ("src/Constants.hpp", "Скорости, дистанции, зум текстур"),
        ("src/LevelData.hpp", "5 уровней, лабиринты, спавны монстров"),
        ("src/Maze.hpp / Maze.cpp", "Сетка, коллизии, BFS, отрисовка пола/стен"),
        ("src/Player.hpp / Player.cpp", "Игрок: ввод, движение, анимация"),
        ("src/PlayerSheet.hpp / .cpp", "Загрузка и нарезка player_sheet.png"),
        ("src/Monster.hpp / Monster.cpp", "ИИ и анимация монстров"),
        ("src/MonsterPinkSheet.hpp / .cpp", "Загрузка monster_pink_sheet.png"),
        ("src/ExitPortal.hpp / .cpp", "Анимированный выход E"),
        ("src/PortalSheet.hpp / .cpp", "Загрузка portal_sheet.png"),
        ("src/MenuChase.hpp / .cpp", "Погоня в главном меню"),
        ("src/Art.hpp / Art.cpp", "Загрузка всех текстур и fallback-кругов"),
        ("src/Utf.hpp", "UTF-8 → sf::String для русского текста"),
        ("TheMazeRunner.vcxproj", "Список .cpp, линковка SFML, PostBuild копирование"),
        ("sfml_config.props", "Путь SFML_DIR, суффиксы lib/dll"),
        ("SETUP.bat", "Установка SFML из zip"),
        ("assets/*", "Графика и звук"),
        ("scripts/extract_*.py", "Пересборка спрайт-листов из исходников"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Файл"
    hdr[1].text = "Назначение"
    for path, desc in rows:
        r = table.add_row().cells
        r[0].text = path
        r[1].text = desc

    h1(doc, "12. Поля класса Game (что хранится в памяти)")
    bullet(doc, "window_ — окно SFML (полный экран)")
    bullet(doc, "state_ — текущее GameState")
    bullet(doc, "levels_, currentLevel_, menuSelectedLevel_ — данные и выбор уровня")
    bullet(doc, "maze_, player_, monsters_, exitPortal_, menuChase_ — игровые объекты")
    bullet(doc, "art_ — ArtPack со всеми текстурами")
    bullet(doc, "levelMusic_, menuMusic_, defeatSound_, victorySound_ — аудио SFML")
    bullet(doc, "timeLeft_ — оставшееся время уровня")
    bullet(doc, "catchingMonster_ — индекс монстра, который «съел» игрока")
    bullet(doc, "titleText_, hudText_, centerText_, hintText_, menuItemText_ — надписи на экране")

    h1(doc, "13. Анимации монстра (MonsterAnimState)")
    bullet(doc, "Run — обычный бег (6 кадров ROW_RUN), стоит — кадр 0")
    bullet(doc, "Transform — 5 кадров превращения в злого при начале погони")
    bullet(doc, "Evil — злой вид, рот открывается по дистанции (колонки 1–4)")
    bullet(doc, "Kill — 4 кадра поимки игрока (ROW_KILL)")
    bullet(doc, "KillDone — анимация завершена, кадр зафиксирован")

    h1(doc, "14. Анимации игрока (PlayerAnimState)")
    bullet(doc, "Idle — ROW_IDLE, направление = lastFacing_")
    bullet(doc, "Run — чередование ROW_RUN_A и ROW_RUN_B")
    bullet(doc, "Dying — 4 кадра: FALL → DEAD_BACK → DEAD_DOWN_A → DEAD_DOWN_B")
    bullet(doc, "Dead — последний кадр смерти до возврата в меню")

    h1(doc, "15. Логика победы и поражения (пошагово)")
    p(doc, "Поражение от монстра: Playing → circlesOverlap → playDefeatSound → "
          "startKill у монстра → Caught → killFinished → player startDeath → Dying → "
          "deathFinished → GameOver.")
    p(doc, "Поражение от времени: timeLeft_ <= 0 → playDefeatSound → Dying → GameOver "
          "(без Caught).")
    p(doc, "Победа: расстояние до exitPos < tileSize×0.45 → playVictorySound → "
          "LevelComplete или GameComplete.")

    h1(doc, "16. Конфигурация Visual Studio (vcxproj)")
    p(doc, "ClCompile — все .cpp из src/. ClInclude — заголовки. "
          "AdditionalDependencies: sfml-graphics, window, audio, system. "
          "PostBuildEvent: копирует DLL SFML + openal32.dll + папку assets в OutDir.")

    h1(doc, "17. Зависимости SFML в проекте")
    bullet(doc, "sfml-graphics — окно, спрайты, текст, примитивы")
    bullet(doc, "sfml-window — события клавиатуры, VideoMode")
    bullet(doc, "sfml-audio — Music и Sound")
    bullet(doc, "sfml-system — Clock, Vector2f, String")

    p(doc, "— Конец документа —")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
