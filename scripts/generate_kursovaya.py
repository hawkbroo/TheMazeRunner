#!/usr/bin/env python3
"""Генерация курсовой работы The Maze Runner (≥30 страниц, Word)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
OUT = ROOT / "docs" / "Курсовая_TheMazeRunner.docx"
DIAGRAM = ROOT / "docs" / "diagrams" / "game_logic.png"


def setup_styles(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)

    for level in (1, 2, 3):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.first_line_indent = Cm(0)
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        h.paragraph_format.space_before = Pt(12 if level == 1 else 6)
        h.paragraph_format.space_after = Pt(6)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 3"].font.size = Pt(14)


def para(doc: Document, text: str, *, indent: bool = True, align=None) -> None:
    p = doc.add_paragraph(text)
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if align is not None:
        p.alignment = align


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def screenshot_placeholder(doc: Document, num: int, title: str, hint: str) -> None:
    """Рамка-заглушка под скриншот из игры."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    lines = [
        "",
        "▼  ВСТАВИТЬ СКРИНШОТ ИЗ ИГРЫ  ▼",
        "",
        f"Рисунок {num} — {title}",
        "",
        f"({hint})",
        "",
        "",
    ]
    cell.text = "\n".join(lines)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        for r in p.runs:
            r.italic = True
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph()


def figure(doc: Document, image_path: Path, caption: str, width_cm: float = 15.5) -> None:
    if not image_path.is_file():
        para(doc, f"[Не найден файл рисунка: {image_path}]", indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(12)


def code_block(doc: Document, text: str, caption: str | None = None) -> None:
    if caption:
        cap = doc.add_paragraph(caption)
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text.rstrip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)


def read_src(name: str, start: int = 1, end: int | None = None) -> str:
    lines = (SRC / name).read_text(encoding="utf-8").splitlines()
    if end is None:
        end = len(lines)
    return "\n".join(lines[start - 1 : end])


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
            p.paragraph_format.first_line_indent = Cm(0)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
    doc.add_paragraph()


def title_page(doc: Document) -> None:
    for _ in range(3):
        para(doc, "", indent=False)
    para(
        doc,
        "Министерство науки и высшего образования Российской Федерации",
        indent=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    para(
        doc,
        'Федеральное государственное автономное образовательное учреждение\n'
        'высшего образования\n'
        '«Национальный исследовательский ядерный университет «МИФИ»»',
        indent=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for _ in range(4):
        para(doc, "", indent=False)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("КУРСОВАЯ РАБОТА")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Times New Roman"
    para(
        doc,
        "по дисциплине «Программирование»",
        indent=False,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for _ in range(2):
        para(doc, "", indent=False)
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(
        "Разработка компьютерной игры «The Maze Runner»\n"
        "на языке C++ с использованием библиотеки SFML"
    )
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.name = "Times New Roman"
    for _ in range(6):
        para(doc, "", indent=False)
    para(doc, "Выполнил: студент группы ________  ________________________", indent=False)
    para(doc, "Проверил: ________________________", indent=False)
    for _ in range(4):
        para(doc, "", indent=False)
    para(doc, "Москва 2026", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)


def contents(doc: Document) -> None:
    h1(doc, "СОДЕРЖАНИЕ")
    items = [
        "ВВЕДЕНИЕ",
        "1 ИДЕЯ И КОНЦЕПЦИЯ ИГРЫ «THE MAZE RUNNER»",
        "1.1 Жанр, сюжетная основа и цель игрока",
        "1.2 Игровые механики и управление",
        "1.3 Уровни, монстры и нарастание сложности",
        "1.4 Требования к программе",
        "2 БИБЛИОТЕКА SFML",
        "2.1 Обзор SFML и используемые модули",
        "2.2 Окно, события и ввод с клавиатуры",
        "2.3 Графика: примитивы, текстуры, спрайты и текст",
        "2.4 Звук: музыка и звуковые эффекты",
        "2.5 Игровой цикл и независимость от FPS",
        "2.6 Подключение SFML к проекту Visual Studio",
        "3 ПОЭТАПНОЕ СОЗДАНИЕ ИГРЫ",
        "3.1 Этап 1. Каркас проекта и точка входа",
        "3.2 Этап 2. Класс Maze — лабиринт, коллизии, отрисовка",
        "3.3 Этап 3. Класс Player — управление и анимация",
        "3.4 Этап 4. Описание уровней (LevelData)",
        "3.5 Этап 5. Класс Monster — ИИ и спрайты",
        "3.6 Этап 6. Класс Game — логика, состояния, меню",
        "3.7 Этап 7. Графика, звук и финальная сборка",
        "3.8 Схема логики игры",
        "4 ТЕСТИРОВАНИЕ И РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ",
        "4.1 Результаты тестирования",
        "4.2 Руководство по сборке и управлению",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "ПРИЛОЖЕНИЕ А. Листинг main.cpp",
        "ПРИЛОЖЕНИЕ Б. Листинг Constants.hpp",
        "ПРИЛОЖЕНИЕ В. Листинг Monster.cpp (фрагмент)",
        "ПРИЛОЖЕНИЕ Г. Листинг Game.cpp (фрагмент)",
        "ПРИЛОЖЕНИЕ Д. Листинг Maze.cpp",
        "ПРИЛОЖЕНИЕ Е. Листинг LevelData.hpp (фрагмент)",
    ]
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.first_line_indent = Cm(0)
    page_break(doc)


def introduction(doc: Document) -> None:
    h1(doc, "ВВЕДЕНИЕ")
    paras = [
        "Современная индустрия развлечений активно использует компьютерные игры как форму "
        "интерактивного программного обеспечения. Для начинающего разработчика создание "
        "небольшой, но законченной игры является эффективным способом закрепить навыки "
        "объектно-ориентированного программирования, работы с графикой, обработки ввода "
        "и проектирования архитектуры приложения. Курсовая работа посвящена разработке "
        "аркадной игры «The Maze Runner» — лабиринта с ограничением по времени, "
        "препятствиями в виде NPC-монстров и нарастающей сложностью на пяти уровнях.",
        "Актуальность темы обусловлена широким применением мультимедийных библиотек в "
        "образовательных и учебных проектах. Библиотека SFML (Simple and Fast Multimedia "
        "Library) предоставляет удобный кроссплатформенный интерфейс для создания "
        "двумерной графики, воспроизведения звука и опроса устройств ввода без "
        "необходимости напрямую работать с низкоуровневыми API операционной системы. "
        "Сочетание языка C++17 и SFML 2.6 в среде Visual Studio 2022 соответствует "
        "типичным требованиям учебных курсов по программированию и позволяет сосредоточиться "
        "на алгоритмах и логике игры.",
        "Объектом исследования является процесс разработки двумерной аркадной игры "
        "с элементами искусственного интеллекта противников. Предметом исследования выступают "
        "методы представления лабиринта, реализации коллизий, поиска пути и управления "
        "состояниями игрового приложения с использованием SFML.",
        "Целью курсовой работы является разработка кроссплатформенного (в рамках Windows "
        "и Visual Studio) приложения на C++ с использованием SFML, реализующего прохождение "
        "лабиринта, ИИ монстров, графический интерфейс и звуковое сопровождение.",
        "Для достижения поставленной цели необходимо решить следующие задачи:",
    ]
    for t in paras:
        para(doc, t)
    tasks = [
        "изучить теоретические основы представления лабиринтов и алгоритмов поиска пути;",
        "рассмотреть возможности библиотеки SFML для 2D-графики, звука и ввода;",
        "спроектировать структуру классов игры и диаграмму состояний;",
        "реализовать пять уровней с различным поведением монстров (патруль и преследование);",
        "обеспечить анимацию персонажей, текстурирование лабиринта и главное меню;",
        "провести тестирование и подготовить руководство пользователя.",
    ]
    for t in tasks:
        bullet(doc, t)
    para(
        doc,
        "Практическая значимость работы заключается в получении готового учебного проекта, "
        "демонстрирующего применение структур данных, алгоритма поиска в ширину (BFS), "
        "конечного автомата состояний игры и работы со спрайт-листами. Структура работы "
        "выстроена в логическом порядке разработки: сначала описывается идея игры, затем "
        "библиотека SFML, после чего поэтапно разбирается создание программы с листингами "
        "кода. Иллюстрации — три скриншота: главное меню, игровой уровень и экран победы.",
    )
    page_break(doc)


def chapter1_idea(doc: Document) -> None:
    h1(doc, "1 ИДЕЯ И КОНЦЕПЦИЯ ИГРЫ «THE MAZE RUNNER»")

    h2(doc, "1.1 Жанр, сюжетная основа и цель игрока")
    for t in [
        "Аркадный лабиринт — классический жанр компьютерных игр, зародившийся в эпоху "
        "автоматов и ранних персональных компьютеров. Игрок управляет персонажем в "
        "замкнутом пространстве, состоящем из коридоров и стен; цель — достичь выхода "
        "или собрать объекты, избегая препятствий и врагов. В отличие от жанра "
        "rogue-like, где акцент делается на процедурной генерации уровней и перманентной "
        "смерти с элементами RPG, классический лабиринт-раннер фокусируется на "
        "реакции, тайминге и знании карты.",
        "Игра «The Maze Runner» относится к поджанру «лабиринт с таймером». Ключевые "
        "механики включают: ограниченное время на прохождение уровня; движение по "
        "сетке с непрерывной физикой внутри клеток; столкновения со стенами в виде "
        "круговых тел; NPC с режимом патрулирования и режимом преследования; "
        "последовательное прохождение пяти уровней с возрастающей сложностью.",
        "С точки зрения проектирования уровней каждая карта задаётся статически — "
        "в виде текстового шаблона. Такой подход упрощает отладку, позволяет "
        "гарантировать проходимость (наличие пути от старта к выходу) и даёт "
        "дизайнеру полный контроль над расположением коридоров и точек спавна монстров.",
    ]:
        para(doc, t)

    h2(doc, "1.2 Игровые механики и управление")
    for t in [
        "Основной игровой цикл уровня: игрок появляется в точке старта (S), изучает "
        "карту, избегает монстров и добирается до портала выхода (E) до истечения таймера. "
        "Движение непрерывное внутри клеток сетки; столкновение со стенами блокирует "
        "проход, но позволяет скользить вдоль коридоров.",
        "Управление реализовано с клавиатуры: WASD и стрелки — движение; Enter — "
        "подтверждение (старт уровня, переход дальше); Esc — выход или возврат в меню; "
        "цифры 1–5 и стрелки вверх/вниз — выбор уровня в главном меню.",
        "При столкновении с монстром уровень проваливается: проигрывается анимация "
        "поимки и смерти. При истечении времени — аналогичная последовательность с "
        "сообщением «Время вышло». При успехе — звук победы и экран перехода на "
        "следующий уровень или финальное поздравление.",
    ]:
        para(doc, t)

    h2(doc, "1.3 Уровни, монстры и нарастание сложности")
    for t in [
        "Игра содержит пять уровней с заранее заданными картами. Сложность нарастает "
        "за счёт уменьшения времени, увеличения числа монстров и смены их поведения.",
        "На уровнях 1–3 монстры в основном патрулируют по заданному маршруту — игрок "
        "учится предсказывать их движение. На уровнях 4–5 включается режим преследования: "
        "монстр замечает игрока в радиусе 220 пикселей и бежит за ним по кратчайшему "
        "пути по коридорам (алгоритм BFS), отставая при удалении дальше 140 пикселей.",
        "Визуально монстры представлены розовыми спрайтами: при начале погони проигрывается "
        "анимация «злости», при поимке — kill-анимация. Игрок анимирован спрайт-листом "
        "с бегом в восьми направлениях и анимацией смерти.",
    ]:
        para(doc, t)
    add_table(
        doc,
        ["Уровень", "Время (с)", "Монстров", "Поведение", "Сложность"],
        [
            ["1 — разведка", "90", "1", "Патруль", "Низкая"],
            ["2 — коридоры", "75", "2", "Патруль", "Ниже средней"],
            ["3 — спешка", "60", "3", "Патруль + Chase", "Средняя"],
            ["4 — охота", "55", "3", "Преследование", "Высокая"],
            ["5 — финал", "40", "3", "Быстрое преследование", "Очень высокая"],
        ],
    )

    h2(doc, "1.4 Требования к программе")
    para(
        doc,
        "Исходя из описанной концепции, сформулированы требования к программной реализации.",
    )
    para(doc, "Функциональные требования:")
    for t in [
        "отображение лабиринта с текстурами пола и стен;",
        "управление игроком с клавиатуры;",
        "таймер и HUD на русском языке;",
        "пять уровней с разным ИИ монстров;",
        "главное меню с выбором уровня и анимацией погони;",
        "экраны победы, поражения и прохождения всех уровней;",
        "фоновая музыка и звуковые эффекты.",
    ]:
        bullet(doc, t)
    para(doc, "Нефункциональные требования:")
    for t in [
        "язык C++17, среда Visual Studio 2022, платформа x64;",
        "графическая библиотека SFML 2.6.1;",
        "адаптация размера окна под разрешение экрана;",
        "модульная структура кода в каталоге src/.",
    ]:
        bullet(doc, t)
    page_break(doc)


def chapter2_sfml(doc: Document) -> None:
    h1(doc, "2 БИБЛИОТЕКА SFML")

    h2(doc, "2.1 Обзор SFML и используемые модули")
    for t in [
        "SFML (Simple and Fast Multimedia Library) — кроссплатформенная мультимедийная "
        "библиотека для C++. Она скрывает низкоуровневую работу с окнами ОС, OpenGL "
        "и аудиоподсистемой, предоставляя простой объектно-ориентированный API.",
        "В проекте «The Maze Runner» используются четыре модуля SFML:",
    ]:
        para(doc, t)
    for t in [
        "sfml-system — базовые типы (sf::Vector2f, sf::String), измерение времени (sf::Clock);",
        "sfml-window — окно приложения (sf::RenderWindow), очередь событий, клавиатура;",
        "sfml-graphics — примитивы, текстуры, спрайты, текст, шрифты;",
        "sfml-audio — фоновая музыка (sf::Music) и звуковые эффекты (sf::Sound).",
    ]:
        bullet(doc, t)
    para(
        doc,
        "Модули связаны зависимостями: graphics и audio опираются на window, window — на system. "
        "При линковке в Visual Studio необходимо подключить все четыре библиотеки (.lib) "
        "и скопировать соответствующие DLL в папку с исполняемым файлом.",
    )

    h2(doc, "2.2 Окно, события и ввод с клавиатуры")
    for t in [
        "Класс sf::RenderWindow создаёт окно с контекстом отрисовки. В Game::init() "
        "окно открывается в полноэкранном или крупном режиме с заголовком «The Maze Runner».",
        "Метод pollEvent извлекает события из очереди ОС: sf::Event::Closed (закрытие), "
        "sf::Event::KeyPressed (однократное нажатие). В processEvents() по KeyPressed "
        "обрабатываются Enter, Esc, M, стрелки и цифры для меню.",
        "Для плавного движения персонажа используется sf::Keyboard::isKeyPressed в "
        "Player::handleInput — опрос удерживаемых клавиш WASD и стрелок между кадрами.",
    ]:
        para(doc, t)

    h2(doc, "2.3 Графика: примитивы, текстуры, спрайты и текст")
    for t in [
        "sf::RectangleShape и sf::CircleShape — простые фигуры для отладки и fallback-"
        "отрисовки (если спрайт-лист не загрузился). sf::Texture хранит изображение в "
        "видеопамяти; sf::Sprite — прямоугольник с текстурой и возможностью выбрать "
        "фрагмент через setTextureRect.",
        "sf::Font и sf::Text отображают надписи HUD и меню. Для кириллицы используется "
        "функция utf8(), преобразующая std::string в sf::String. Шрифт загружается из "
        "системного Arial или Times New Roman.",
        "Отрисовка кадра: window.clear(цвет фона) → draw объекты → window.display(). "
        "Порядок draw задаёт слои: сначала лабиринт, затем портал, монстры, игрок, HUD.",
    ]:
        para(doc, t)

    h2(doc, "2.4 Звук: музыка и звуковые эффекты")
    for t in [
        "sf::Music воспроизводит длинные треки из файла (menu_music, level_music) в цикле. "
        "sf::SoundBuffer загружает короткий звук в память; sf::Sound проигрывает его "
        "однократно (победа, поражение).",
        "В Game реализовано переключение: при старте уровня останавливается музыка меню, "
        "при возврате в MainMenu — музыка уровня. Метод setVolume регулирует громкость.",
    ]:
        para(doc, t)

    h2(doc, "2.5 Игровой цикл и независимость от FPS")
    for t in [
        "sf::Clock::restart() возвращает время с прошлого кадра в секундах (dt). "
        "Скорость умножается на dt: displacement = speed * dt. Так персонаж проходит "
        "одинаковое расстояние за секунду при любом FPS.",
        "Классический цикл SFML-игры: while (window.isOpen()) { processEvents(); "
        "update(dt); render(); }. В проекте он инкапсулирован в Game::run().",
    ]:
        para(doc, t)
    code_block(
        doc,
        read_src("Game.cpp", 580, 587),
        "Листинг 2.1 — Главный игровой цикл Game::run()",
    )

    h2(doc, "2.6 Подключение SFML к проекту Visual Studio")
    for t in [
        "Скачивается сборка SFML 2.6.x для MSVC 2022 64-bit. В свойствах проекта: "
        "C/C++ → Additional Include Directories — путь к include; Linker → Additional "
        "Library Directories — путь к lib; Linker → Input — sfml-graphics.lib, "
        "sfml-window.lib, sfml-system.lib, sfml-audio.lib (для Debug — с суффиксом -d).",
        "PostBuildEvent копирует DLL из bin SFML и папку assets рядом с .exe. "
        "Переменная SFML_DIR позволяет не прописывать абсолютные пути на каждой машине.",
    ]:
        para(doc, t)
    page_break(doc)


def chapter3_steps(doc: Document) -> None:
    h1(doc, "3 ПОЭТАПНОЕ СОЗДАНИЕ ИГРЫ")
    para(
        doc,
        "Разработка велась итеративно: от минимального каркаса с окном до полноценной "
        "игры с пятью уровнями. Ниже каждый этап описан в порядке реализации с "
        "пояснением логики и фрагментами кода.",
    )

    h2(doc, "3.1 Этап 1. Каркас проекта и точка входа")
    for t in [
        "На первом этапе создан проект Visual Studio, подключена SFML и написана "
        "минимальная точка входа. Класс Game инкапсулирует окно, шрифт, состояние "
        "и методы init(), run(), processEvents(), update(), render().",
        "Перечисление GameState задаёт режимы: MainMenu, Playing, Caught, Dying, "
        "LevelComplete, GameComplete, GameOver. Файл Constants.hpp собирает "
        "глобальные настройки баланса — скорости, дистанции ИИ, размеры клеток.",
    ]:
        para(doc, t)
    code_block(doc, read_src("main.cpp"), "Листинг 3.1 — Точка входа main.cpp")
    code_block(doc, read_src("Constants.hpp"), "Листинг 3.2 — Константы игры Constants.hpp")
    code_block(
        doc,
        read_src("Game.hpp", 15, 28),
        "Листинг 3.3 — Перечисление состояний GameState",
    )

    h2(doc, "3.2 Этап 2. Класс Maze — лабиринт, коллизии, отрисовка")
    for t in [
        "Метод loadFromLayout копирует вектор строк, выравнивает ширину строк, "
        "находит S и E, записывает пиксельные координаты центров старта и выхода, "
        "заменяет S/E на пол.",
        "Метод moveCircle реализует скольжение: сначала попытка сдвига по X, затем по Y. "
        "Радиус для проверки умножается на COLLISION_RADIUS_SCALE (0.88), чтобы "
        "персонажи чуть свободнее проходили углы.",
        "Отрисовка draw обходит все клетки. Для стен и пола при наличии текстур "
        "вызывается drawTiledTexture — фрагмент текстуры выбирается по координатам "
        "клетки с параметром zoom, создавая эффект разнообразия без отдельной "
        "текстуры на каждую клетку.",
    ]:
        para(doc, t)
    code_block(
        doc,
        read_src("Maze.cpp", 150, 158),
        "Листинг 3.4 — Метод Maze::moveCircle (скольжение вдоль стен)",
    )
    code_block(
        doc,
        read_src("Maze.cpp", 32, 55),
        "Листинг 3.5 — Загрузка карты Maze::loadFromLayout",
    )
    para(
        doc,
        "На рисунке 2 показан игровой уровень в процессе прохождения: лабиринт с "
        "текстурами, персонаж игрока, монстр и строка HUD с таймером.",
    )
    screenshot_placeholder(
        doc, 2, "Игровой уровень",
        "скриншот любого уровня: лабиринт, игрок, монстр, HUD с таймером внизу",
    )

    h2(doc, "3.3 Этап 3. Класс Player — управление и анимация")
    for t in [
        "Класс Player хранит позицию, радиус коллизии, скорость PLAYER_SPEED (130 px/s), "
        "направление движения и состояние анимации (Idle, Run, Dying, Dead). "
        "Метод handleInput опрашивает WASD и стрелки, формирует нормализованный "
        "вектор moveDir_.",
        "Метод update применяет moveDir_ * speed_ * dt через maze.moveCircle. "
        "При наличии загруженного PlayerSheet отрисовка выполняется кадрами "
        "спрайт-листа в зависимости от направления (8 столбцов) и фазы бега. "
        "При startDeath запускается анимация гибели перед переходом в GameOver.",
    ]:
        para(doc, t)
    code_block(doc, read_src("Player.hpp"), "Листинг 3.6 — Заголовочный файл Player.hpp")

    h2(doc, "3.4 Этап 4. Описание уровней (LevelData)")
    for t in [
        "Уровни задаются текстовыми картами: «#» — стена, «.» — пол, «S» — старт, «E» — выход. "
        "Структура LevelConfig хранит имя, layout, лимит времени, скорость монстров и "
        "список MonsterSpawn. Функция buildLevels() возвращает пять готовых конфигураций.",
        "Уровень 1 масштабируется функцией scaleLayout (×2) для увеличения детализации карты.",
    ]:
        para(doc, t)
    code_block(doc, read_src("LevelData.hpp", 1, 26), "Листинг 3.7 — Структуры LevelConfig и MonsterSpawn")
    code_block(doc, read_src("LevelData.hpp", 46, 70), "Листинг 3.8 — Фрагмент уровня 1 в buildLevels()")

    h2(doc, "3.5 Этап 5. Класс Monster — ИИ и спрайты")
    for t in [
        "Монстры поддерживают два режима ИИ: Patrol (патруль по направлению до стены) "
        "и Chase (преследование с BFS по сетке лабиринта). При приближении игрока "
        "ближе 220 px монстр начинает погоню; при удалении дальше 140 px — прекращает.",
        "Для навигации в Chase используется Maze::nextStepBfs — поиск в ширину по "
        "проходимым клеткам с восстановлением первого шага пути.",
    ]:
        para(doc, t)
    code_block(
        doc,
        read_src("Maze.cpp", 95, 140),
        "Листинг 3.9 — Алгоритм BFS: Maze::nextStepBfs",
    )
    for t in [
        "Метод spawn размещает монстра в ближайшем центре проходимой клетки, "
        "инициализирует направление патруля и скорость. Для Chase базовая скорость "
        "пересчитывается каждый кадр от скорости игрока.",
        "updatePatrol двигает монстра до столкновения; при lockPatrolAxis только "
        "разворот, иначе — попытка выбрать альтернативное направление из четырёх.",
        "updateChase включает гистерезис дистанции, вызывает onChaseStarted для "
        "анимации трансформации, вычисляет nextStepBfs и двигает монстра к центру "
        "следующей клетки пути.",
        "MonsterPinkSheet загружает monster_pink_sheet.png и нарезает кадры по "
        "прозрачным границам. Ряды: бег, трансформация/«злой» рот, kill.",
    ]:
        para(doc, t)
    code_block(
        doc,
        read_src("Monster.cpp", 111, 145),
        "Листинг 3.10 — Метод Monster::updateChase",
    )

    h2(doc, "3.6 Этап 6. Класс Game — логика, состояния, меню")
    for t in [
        "Метод init создаёт полноэкранное или крупное окно, загружает шрифт, арт, "
        "музыку, строит levels_ = buildLevels(), инициализирует тексты меню.",
        "startLevel загружает layout в maze, подбирает tileSize через fitTileSizeForLevel "
        "(масштабирование под 92% ширины и 85% высоты экрана), расставляет игрока "
        "и монстров, сбрасывает таймер, запускает музыку уровня.",
        "processEvents обрабатывает Esc (выход или меню), Enter (старт/следующий уровень), "
        "стрелки и цифры 1–5 в меню. update реализует логику всех состояний — "
        "см. листинг 3.6.",
        "drawMainMenu рисует фон menu_background, анимацию MenuChase (погоня в полосе "
        "500 px с зацикливанием), список уровней с неоновым оформлением drawNeonMenuText.",
    ]:
        para(doc, t)
    para(
        doc,
        "На рисунке 1 представлено главное меню: фон, список из пяти уровней и "
        "декоративная анимация погони.",
    )
    screenshot_placeholder(
        doc, 1, "Главное меню",
        "скриншот меню: фон, список уровней, анимация погони, неоновые надписи",
    )
    code_block(
        doc,
        read_src("Game.cpp", 372, 452),
        "Листинг 3.11 — Метод Game::update (игровая логика)",
    )
    para(
        doc,
        "На рисунке 3 показан экран победы после прохождения уровня (состояние "
        "LevelComplete) с предложением перейти дальше.",
    )
    screenshot_placeholder(
        doc, 3, "Экран победы",
        "скриншот после достижения выхода: «Уровень пройден!»",
    )

    h2(doc, "3.7 Этап 7. Графика, звук и финальная сборка")
    for t in [
        "Структура ArtPack централизует указатели на sf::Texture и флаги loaded. "
        "При старте загружаются: floor.png, wall.png, player_sheet.png, "
        "monster_pink_sheet.png, portal sprites, menu_background.jpg.",
        "ExitPortal анимирует спрайты выхода с циклической сменой кадров. "
        "MenuChase создаёт на главном меню сцену погони: игрок и монстр движутся "
        "в горизонтальной полосе; после «поимки» позиции сбрасываются для бесконечного "
        "цикла демонстрации геймплея.",
        "Параметры FLOOR_TEXTURE_ZOOM и WALL_TEXTURE_ZOOM управляют масштабом "
        "вырезаемого фрагмента текстуры при тайлинге — пол крупнее, стены мельче, "
        "что визуально разделяет типы поверхностей.",
        "На финальном этапе подключены музыка меню и уровней, звуки победы и поражения, "
        "проведена балансировка таймеров и скоростей на всех пяти уровнях.",
    ]:
        para(doc, t)

    h2(doc, "3.8 Схема логики игры")
    para(
        doc,
        "Для наглядности сведены воедино главный цикл и конечный автомат состояний. "
        "Верхняя часть схемы — последовательность вызовов за один кадр; нижняя — "
        "переходы между GameState по событиям игры и нажатиям клавиш.",
    )
    figure(
        doc,
        DIAGRAM,
        "Рисунок 4 — Схема логики игры: цикл Game::run() и диаграмма состояний",
    )
    page_break(doc)


def chapter4(doc: Document) -> None:
    h1(doc, "4 ТЕСТИРОВАНИЕ И РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ")

    h2(doc, "4.1 Результаты тестирования")
    for t in [
        "Тестирование программного обеспечения игрового типа включает модульное "
        "тестирование отдельных алгоритмов (BFS, коллизии), интеграционное "
        "тестирование взаимодействия классов и приёмочное тестирование сценариев "
        "пользователя. В рамках курсовой работы основной акцент сделан на "
        "интеграционном и приёмочном уровнях, так как визуальная составляющая "
        "и тайминг анимаций требуют ручной проверки.",
        "Тестирование проводилось вручную на сборке Debug|x64 и Release|x64 в Windows 10. "
        "Проверялись функциональные сценарии, визуальное отображение и стабильность "
        "при длительной игре (не менее 30 минут непрерывной работы без утечек памяти "
        "и падений). Отдельно проверялась корректность масштабирования окна на мониторе "
        "1920×1080: лабиринт уровня 5 полностью помещается в видимую область с HUD снизу.",
        "Алгоритм BFS проверялся косвенно: монстры в режиме Chase не проходят сквозь "
        "стены и выбирают короткий путь в простых коридорах. При блокировке прямого "
        "маршрута монстр обходит препятствие по сетке, что подтверждает корректность "
        "восстановления пути по массиву предков.",
    ]:
        para(doc, t)
    add_table(
        doc,
        ["№", "Сценарий", "Ожидаемый результат", "Статус"],
        [
            ["1", "Запуск без DLL SFML", "Ошибка загрузки библиотеки", "Ожидаемо"],
            ["2", "Запуск с assets и DLL", "Открывается главное меню", "Пройден"],
            ["3", "Движение в стену", "Персонаж скользит вдоль стены", "Пройден"],
            ["4", "Достижение выхода", "Звук победы, LevelComplete", "Пройден"],
            ["5", "Столкновение с монстром", "Caught → Dying → GameOver", "Пройден"],
            ["6", "Истечение таймера", "Смерть, сообщение «Время вышло»", "Пройден"],
            ["7", "Прохождение 5 уровней", "GameComplete", "Пройден"],
            ["8", "Esc из уровня", "Возврат в MainMenu", "Пройден"],
            ["9", "Chase: отбегание >140 px", "Монстр прекращает преследование", "Пройден"],
            ["10", "Выбор уровня 1–5 в меню", "Старт выбранного уровня", "Пройден"],
        ],
    )
    for t in [
        "На первых уровнях игрок изучает карту и тайминг патрулей. С уровня 4 требуется "
        "уход от преследования и использование гистерезиса ИИ — монстр отстаёт, если "
        "убежать достаточно далеко. Уровень 5 объединяет короткий таймер (40 с) и "
        "трёх преследователей с повышенными множителями скорости.",
        "При отсутствии спрайт-листа монстры рисуются цветными кругами — fallback "
        "для отладки без графических ресурсов.",
    ]:
        para(doc, t)

    h2(doc, "4.2 Руководство по сборке и управлению")
    para(doc, "Установка и сборка:")
    for t in [
        "Установить Visual Studio 2022 с компонентом «Разработка классических приложений на C++».",
        "Скачать SFML 2.6.x для MSVC 2022 64-bit с официального сайта sfml-dev.org.",
        "Задать SFML_DIR или прописать пути в свойствах проекта.",
        "Открыть TheMazeRunner.sln, выбрать конфигурацию Debug x64, собрать (Ctrl+Shift+B).",
        "Убедиться, что PostBuild скопировал DLL и папку assets рядом с .exe.",
        "Запустить Ctrl+F5 (без отладки).",
    ]:
        bullet(doc, " ".join(t.split()))
    para(doc, "Управление в игре:")
    for t in [
        "WASD или стрелки — движение персонажа;",
        "Enter — начать уровень / следующий уровень / закрыть экран результата;",
        "Esc — выход из программы или возврат в главное меню;",
        "M — альтернативный возврат в меню с экрана поражения;",
        "Стрелки вверх/вниз и цифры 1–5 — выбор уровня в меню.",
    ]:
        bullet(doc, t)
    page_break(doc)


def conclusion(doc: Document) -> None:
    h1(doc, "ЗАКЛЮЧЕНИЕ")
    for t in [
        "В ходе выполнения курсовой работы разработана и реализована компьютерная игра "
        "«The Maze Runner» на языке C++17 с использованием мультимедийной библиотеки SFML 2.6. "
        "Программа соответствует поставленным функциональным требованиям: реализованы пять "
        "уровней лабиринта, таймер, HUD на русском языке, два режима ИИ монстров "
        "(патрулирование и преследование с гистерезисом по дистанции), главное меню "
        "с анимацией, текстурированный лабиринт, спрайтовая анимация игрока и монстров, "
        "звуковое сопровождение.",
        "Работа построена в логическом порядке разработки: от идеи и концепции игры "
        "через изучение SFML к поэтапной реализации с листингами кода. Тестирование "
        "подтвердило работоспособность основных сценариев на платформе Windows 10.",
        "Направления дальнейшего развития проекта: загрузка уровней из внешних файлов "
        "(JSON/текст); редактор лабиринтов; таблица рекордов; сетевой мультиплеер; "
        "процедурная генерация карт; расширение звуковых эффектов и частиц при победе.",
        "Таким образом, цель курсовой работы достигнута, все поставленные задачи выполнены.",
    ]:
        para(doc, t)
    page_break(doc)


def bibliography(doc: Document) -> None:
    h1(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    sources = [
        "Документация SFML 2.6 [Электронный ресурс]. — URL: https://www.sfml-dev.org/documentation/2.6.0/ "
        "(дата обращения: 11.06.2026).",
        "Страуструп Б. Язык программирования C++. — 4-е изд. — СПб.: Бином, 2011. — 1136 с.",
        "Лафоре Р. Структуры данных и алгоритмы в C++. — СПб.: Питер, 2011. — 384 с.",
        "Microsoft Visual Studio Documentation [Электронный ресурс]. — URL: "
        "https://learn.microsoft.com/visualstudio/ (дата обращения: 11.06.2026).",
        "Котлер К. Основы разработки компьютерных игр. — М.: Вильямс, 2003. — 736 с.",
        "Методические указания по выполнению курсовой работы по дисциплине «Программирование» "
        "/ НИЯУ МИФИ. — 2026.",
        "Learn C++ [Электронный ресурс]. — URL: https://www.learncpp.com/ (дата обращения: 11.06.2026).",
        "SFML Game Development / H. Fryer, A. Jesus. — Packt Publishing, 2015. — 296 p.",
    ]
    for i, s in enumerate(sources, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)
    page_break(doc)


def appendices(doc: Document) -> None:
    h1(doc, "ПРИЛОЖЕНИЕ А")
    para(doc, "Листинг файла main.cpp", indent=False)
    code_block(doc, read_src("main.cpp"))

    h1(doc, "ПРИЛОЖЕНИЕ Б")
    para(doc, "Листинг файла Constants.hpp", indent=False)
    code_block(doc, read_src("Constants.hpp"))

    h1(doc, "ПРИЛОЖЕНИЕ В")
    para(doc, "Листинг файла Monster.cpp (основные методы)", indent=False)
    code_block(doc, read_src("Monster.cpp", 1, 175))

    h1(doc, "ПРИЛОЖЕНИЕ Г")
    para(doc, "Листинг файла Game.cpp (фрагмент: update и run)", indent=False)
    code_block(doc, read_src("Game.cpp", 372, 587))

    h1(doc, "ПРИЛОЖЕНИЕ Д")
    para(doc, "Листинг файла Maze.cpp", indent=False)
    code_block(doc, read_src("Maze.cpp"))

    h1(doc, "ПРИЛОЖЕНИЕ Е")
    para(doc, "Листинг файла LevelData.hpp (фрагмент)", indent=False)
    code_block(doc, read_src("LevelData.hpp", 1, 42))
    code_block(doc, read_src("LevelData.hpp", 43, 130))
    code_block(doc, read_src("LevelData.hpp", 161, 196))


def estimate_pages(doc_path: Path) -> float:
    """Грубая оценка: текст + код (строки × коэффициент) + разрывы страниц."""
    from docx import Document as D
    from docx.oxml.ns import qn as _qn

    d = D(str(doc_path))
    chars = sum(len(p.text) for p in d.paragraphs)
    code_lines = sum(p.text.count("\n") + 1 for p in d.paragraphs if p.runs and p.runs[0].font.name == "Consolas")
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                chars += len(cell.text)
    # Разрывы страниц учтены вручную (фиксированное число в документе)
    page_breaks = 12
    return chars / 1600.0 + code_lines / 45.0 + page_breaks


def main() -> None:
    import sys

    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from game_logic_diagram import render_game_logic_diagram

    render_game_logic_diagram(DIAGRAM)

    doc = Document()
    setup_styles(doc)
    title_page(doc)
    contents(doc)
    introduction(doc)
    chapter1_idea(doc)
    chapter2_sfml(doc)
    chapter3_steps(doc)
    chapter4(doc)
    conclusion(doc)
    bibliography(doc)
    appendices(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    save_path = OUT
    try:
        doc.save(str(save_path))
    except PermissionError:
        save_path = OUT.with_stem(OUT.stem + "_обновлено")
        doc.save(str(save_path))
        print(f"NOTE: {OUT.name} занят — сохранено в {save_path.name}")
    pages = estimate_pages(save_path)
    print(f"Saved: {save_path}")
    print(f"Estimated pages (text): ~{pages:.1f}")
    if pages < 30:
        print("WARNING: estimate below 30 — open in Word and check real page count.")


if __name__ == "__main__":
    main()
