#!/usr/bin/env python3
"""Очень подробный Word-документ про NPC — логика + каждая строка кода."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "TheMazeRunner_NPC_подробно.docx"


def setup(doc: Document) -> None:
    for sec in doc.sections:
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)


def h1(doc: Document, t: str) -> None:
    doc.add_heading(t, level=1)


def h2(doc: Document, t: str) -> None:
    doc.add_heading(t, level=2)


def h3(doc: Document, t: str) -> None:
    doc.add_heading(t, level=3)


def p(doc: Document, t: str) -> None:
    para = doc.add_paragraph(t)
    para.paragraph_format.space_after = Pt(6)


def logic_block(doc: Document, t: str) -> None:
    para = doc.add_paragraph()
    r = para.add_run("▶ Логика: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0, 80, 160)
    para.add_run(t)
    para.paragraph_format.space_after = Pt(8)


def add_lines(doc: Document, rows: list[tuple[str, str]], title: str | None = None) -> None:
    """rows = [(code_line, explanation), ...]"""
    if title:
        h3(doc, title)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Строка кода"
    hdr[2].text = "Что делает и зачем"
    for i, (code, expl) in enumerate(rows, 1):
        c = table.add_row().cells
        c[0].text = str(i)
        c[1].text = code
        c[2].text = expl
    doc.add_paragraph()


# --- данные построчно ---

MONSTER_HPP = [
    ('#pragma once', 'Стандартная защита: файл не подключится дважды.'),
    ('#include <SFML/Graphics.hpp>', 'Графика SFML: спрайты, векторы, отрисовка.'),
    ('#include "LevelData.hpp"', 'Структуры MonsterSpawn, MonsterAI — откуда берутся настройки монстра.'),
    ('#include "Maze.hpp"', 'Лабиринт: движение, BFS, коллизии со стенами.'),
    ('#include "MonsterPinkSheet.hpp"', 'Спрайт-лист розового монстра.'),
    ('enum class MonsterAnimState { Run, Transform, Evil, Kill, KillDone };',
     'Состояния анимации: бег | превращение | злой | поимка | поимка закончена.'),
    ('void spawn(const MonsterSpawn& cfg, float baseSpeed, const Maze& maze);',
     'Поставить монстра на уровень по данным из LevelData.'),
    ('void setPinkSheet(const MonsterPinkSheet* sheet) { pinkSheet_ = sheet; }',
     'Указатель на общий лист текстур (один на всех монстров).'),
    ('void update(float dt, const Maze& maze, sf::Vector2f playerPos, float playerSpeed);',
     'Главный метод каждого кадра: движение + анимация.'),
    ('void updateKill(float dt);', 'Отдельное обновление анимации kill, когда игрок уже пойман.'),
    ('void startKill();', 'Запуск анимации «съел игрока» при столкновении.'),
    ('bool isKilling() / killFinished()', 'Game спрашивает: идёт ли kill и закончился ли.'),
    ('void draw(sf::RenderTarget& target) const;', 'Нарисовать монстра на экране.'),
    ('sf::Vector2f position() / float radius()', 'Для проверки столкновения с игроком в Game.'),
    ('void updatePatrol / updateChase (private)', 'Два режима ИИ — патруль и погоня.'),
    ('void onChaseStarted / onChaseCalmed', 'Реакция на начало/конец погони (анимация злости).'),
    ('void advanceAnim(float dt, bool moving)', 'Переключение кадров спрайта по таймеру.'),
    ('void drawSprite / mouthColFromDist', 'Выбор кадра на листе и отрисовка; рот по дистанции.'),
    ('sf::Vector2f pos_{};', 'Позиция центра монстра в пикселях на карте.'),
    ('float radius_{12.f};', 'Радиус для коллизий и размера на экране.'),
    ('float speed_{60.f};', 'Текущая скорость px/с (меняется у chase).'),
    ('float speedMultiplier_{1.f};', 'Множитель из LevelData (быстрее/медленнее).'),
    ('float playerDist_{9999.f};', 'Дистанция до игрока — для «открытия рта».'),
    ('MonsterAI ai_{MonsterAI::Patrol};', 'Patrol или Chase — тип поведения.'),
    ('int patrolDirX_{1}; int patrolDirY_{0};', 'Направление патруля: (1,0)=вправо, (0,1)=вниз и т.д.'),
    ('bool lockPatrolAxis_{false};', 'true = только разворот на 180°, не сворачивать на перекрёстке.'),
    ('bool chasing_{false};', 'Сейчас преследует игрока (только у Chase AI).'),
    ('bool evil_{false};', 'Уже превратился в злого (после Transform).'),
    ('bool facingLeft_{false};', 'Куда смотрит — для зеркалирования спрайта.'),
    ('const MonsterPinkSheet* pinkSheet_{nullptr};', 'Ссылка на текстуру; без неё — цветной круг.'),
    ('MonsterAnimState animState_{Run};', 'Текущее состояние анимации.'),
    ('int animFrame_{0};', 'Номер кадра в текущем ряду листа.'),
    ('float animTimer_{0.f};', 'Накопленное время с прошлого кадра.'),
    ('bool lastMoving_{false};', 'Двигался ли в этом кадре — для анимации бега.'),
]

PINK_SHEET = [
    ('struct MonsterPinkSheet {', 'Обёртка над PNG-листом 6×3 клетки.'),
    ('sf::Texture texture;', 'Загруженная картинка monster_pink_sheet.png.'),
    ('unsigned cellW, cellH;', 'Ширина/высота одной клетки = размер текстуры / COLS / ROWS.'),
    ('bool loaded{false};', 'Успешно ли загрузился файл.'),
    ('COLS = 6, ROWS = 3', 'Сетка: 6 колонок, 3 ряда анимаций.'),
    ('ROW_RUN = 0', 'Верхний ряд — 6 кадров обычного бега.'),
    ('ROW_TRANSFORM = 1', 'Средний ряд — 5 кадров превращения в злого.'),
    ('ROW_KILL = 2', 'Нижний ряд — 5 кадров поимки (в игре 4).'),
    ('RUN_FRAMES = 6', 'Сколько кадров циклически крутить в Run/Evil.'),
    ('TRANSFORM_FRAMES = 5', 'Длина анимации злости при старте погони.'),
    ('KILL_PLAY_FRAMES = 4', 'Показываем 4 кадра kill (5-й на листе не используем).'),
    ('loadFromFile(path)', 'texture.loadFromFile → делит размер на COLS×ROWS.'),
    ('cellRect(col, row)', 'Возвращает sf::IntRect — прямоугольник кадра в пикселях текстуры.'),
]

MONSTER_CPP_HEADER = [
    ('#include "Monster.hpp"', 'Подключение своего заголовка.'),
    ('#include <algorithm>', 'std::min для кадров смерти/kill.'),
    ('#include <cmath>', 'sqrt для дистанций и нормализации векторов.'),
    ('#include "Constants.hpp"', 'CHASE_START_DISTANCE, MONSTER_CHASE_FACTOR и др.'),
    ('constexpr float RUN_FRAME_TIME = 0.09f;', 'Длительность одного кадра бега — ~11 кадров/с.'),
    ('constexpr float TRANSFORM_FRAME_TIME = 0.14f;', 'Кадр превращения дольше — заметнее анимация.'),
    ('constexpr float KILL_FRAME_TIMES[] = {...}', 'Разная длительность каждого кадра поимки.'),
    ('sf::Vector2f dirFromPatrol(int px, int py)', 'Превращает (-1/0/1, -1/0/1) в единичный вектор направления.'),
    ('  d = (px, py); d /= len;', 'Например (1,0)→вправо, (0,-1)→вверх. Длина=1 для равномерной скорости.'),
]

SPAWN_LINES = [
    ('pos_ = maze.nearestFloorCenter(cfg.gridX, cfg.gridY);',
     'Ставим монстра в центр клетки (gridX, gridY). Если там стена — BFS ищет ближайший проход.'),
    ('patrolDirX_ = cfg.patrolDirX;', 'Копируем направление из LevelData, напр. 1 = вправо.'),
    ('patrolDirY_ = cfg.patrolDirY;', '0 = не ходит по вертикали при горизонтальном патруле.'),
    ('lockPatrolAxis_ = cfg.lockPatrolAxis;', 'Флаг «не сходить с линии коридора».'),
    ('ai_ = cfg.ai;', 'Patrol или Chase.'),
    ('chasing_ = false;', 'Пока никого не преследует.'),
    ('evil_ = false;', 'Ещё не злой (розовый обычный вид).'),
    ('facingLeft_ = patrolDirX_ < 0;', 'Если патруль влево — спрайт зеркалим.'),
    ('speedMultiplier_ = cfg.speedMultiplier;', '1.05 = на 5% быстрее базовой скорости.'),
    ('animState_ = Run; animFrame_ = 0; animTimer_ = 0;', 'Стартуем с первого кадра бега.'),
    ('playerDist_ = 9999.f;', 'Игрок «очень далеко» — рот закрыт.'),
    ('if (ai_ == Patrol) speed_ = baseSpeed * mult;', 'Патруль: скорость уровня × множитель.'),
    ('else speed_ = 45.f * mult;', 'Chase в режиме патруля: фикс. 45 × множитель (потом пересчитается).'),
]

PATROL_LINES = [
    ('wish = dirFromPatrol(...) * speed_ * dt', 'На сколько пикселей сдвинуться за этот кадр.'),
    ('before = pos_', 'Запоминаем позицию до движения.'),
    ('maze.moveCircle(pos_, wish, radius_)', 'Двигаем с коллизией стен (отдельно X и Y).'),
    ('if (сдвиг² < 0.01) — упёрлись', 'Почти не сдвинулись = стена впереди.'),
    ('if (lockPatrolAxis_) разворот dir *= -1', 'Только разворот: (1,0)→(-1,0).'),
    ('else перебор 4 направлений', 'Пробуем вправо/влево/вниз/вверх — первый проходимый.'),
    ('если никуда — разворот', 'Как крайний случай в тупике.'),
]

CHASE_LINES = [
    ('speed_ = playerSpeed * 0.82 * mult', 'В погоне скорость привязана к игроку, но чуть меньше (82%).'),
    ('dist = расстояние до playerPos', 'Евклидова дистанция в пикселях.'),
    ('playerDist_ = dist', 'Сохраняем для анимации рта.'),
    ('wasChasing = chasing_', 'Был ли в погоне на прошлом шаге — для детекта «только что начал».'),
    ('if (!chasing && dist < 220) chasing_ = true', 'Включить погоню, если игрок ближе 220 px.'),
    ('if (chasing && dist > 140) chasing_ = false', 'Выключить, если отбежал дальше 140 px (гистерезис).'),
    ('if (!wasChasing && chasing_) onChaseStarted()', 'Первый кадр погони → анимация Transform.'),
    ('if (!chasing_) { onChaseCalmed(); updatePatrol(); return; }', 'Далеко — ведём себя как патруль.'),
    ('if (animState == Transform) return;', 'Во время превращения стоим на месте.'),
    ('fromG = pixelToGrid(pos_)', 'В какой клетке сетки стоит монстр.'),
    ('toG = pixelToGrid(playerPos)', 'В какой клетке игрок.'),
    ('nextG = maze.nextStepBfs(fromG, toG)', 'BFS: следующая клетка на кратчайшем пути.'),
    ('if (nextG == fromG) return;', 'Пути нет или уже на месте — не двигаемся.'),
    ('target = gridCenter(nextG)', 'Центр следующей клетки в пикселях.'),
    ('dir = normalize(target - pos_)', 'Единичный вектор к цели.'),
    ('moveCircle(pos_, dir * speed_ * dt, radius_)', 'Шаг к игроку с учётом стен.'),
]

UPDATE_LINES = [
    ('if (Kill || KillDone) { advanceAnim; return; }', 'Во время поимки только анимация, без движения.'),
    ('before = pos_', 'Для определения, двигался ли монстр.'),
    ('if (Patrol) updatePatrol + dist до игрока', 'Патрульный только считает дистанцию (рот не нужен).'),
    ('else updateChase(...)', 'Chase — вся логика погони.'),
    ('lastMoving_ = (delta² > 0.25)', 'Сдвинулся больше чем на 0.5 px — считаем «бежит».'),
    ('facingLeft_ по delta.x или patrolDirX', 'Куда рисовать спрайт.'),
    ('advanceAnim(dt, lastMoving_)', 'Обновить номер кадра анимации.'),
]

ADVANCE_ANIM = [
    ('if (!pinkSheet_) return;', 'Без текстуры анимации нет.'),
    ('animTimer_ += dt', 'Накапливаем время кадра.'),
    ('Run: if (!moving) frame=0', 'Стоит — первый кадр бега (idle-поза).'),
    ('Run: if timer>=0.09 → frame++ % 6', 'Цикл 6 кадров бега.'),
    ('Evil: как Run, но только если moving', 'Злой бежит — кадры бега в фоне для wobble рта.'),
    ('Transform: timer>=0.14 → frame++', 'По одному кадру превращения.'),
    ('if frame>=5 → evil_=true, state=Evil', 'Превращение закончено — теперь злой.'),
    ('Kill: KILL_FRAME_TIMES[frame] → frame++', 'Разная скорость кадров поимки.'),
    ('if frame>=4 → KillDone', 'Анимация поимки завершена.'),
]

DRAW_LINES = [
    ('if (pinkSheet) drawSprite; else CircleShape', 'Спрайт или оранжевый/красный круг-заглушка.'),
    ('Run → row=ROW_RUN, col=animFrame_', 'Кадр бега из верхнего ряда.'),
    ('Transform → row=ROW_TRANSFORM, col=frame', 'Средний ряд — превращение.'),
    ('Evil → row=ROW_TRANSFORM, col=mouthColFromDist', 'Тот же ряд, но колонка по дистанции 1–4.'),
    ('if col<1 col=1', 'Минимум «слегка злой».'),
    ('if moving && frame%2 → col++', 'Лёгкое «дрожание» рта при беге.'),
    ('Kill → row=ROW_KILL', 'Нижний ряд поимки.'),
    ('setTextureRect(cellRect(col,row))', 'Вырезаем нужный прямоугольник из листа.'),
    ('displayH = radius * 4', 'Высота спрайта на экране.'),
    ('setOrigin(0.5, 0.55)', 'Точка привязки — чуть ниже центра (ноги).'),
    ('setScale(facingLeft?-scale:scale, scale)', 'Отрицательный scaleX = отражение влево.'),
]

GAME_NPC = [
    ('for (spawn : lvl.monsters) { Monster m; m.spawn(...); }',
     'На старте уровня создаём монстров по списку из LevelData.'),
    ('m.setPinkSheet(&art_.monsterPinkSheet)', 'Всем один и тот же PNG-лист.'),
    ('monsters_.push_back(m)', 'Храним в векторе — по одному объекту на монстра.'),
    ('monster.update(dt, maze_, player_.position(), player_.speed())',
     'Каждый кадр Playing: двигаем всех монстров.'),
    ('circlesOverlap(player, monster)', 'Столкновение двух кругов: dist < r1+r2.'),
    ('catchingMonster_ = i; startKill(); state=Caught', 'Запоминаем кто поймал, kill-анимация.'),
    ('updateKill пока не killFinished → player.startDeath', 'После kill — смерть игрока.'),
]

BFS_LINES = [
    ('if (from == to) return from', 'Уже на клетке игрока.'),
    ('prev[] = -1; queue; push(from)', 'Классический BFS по клеткам пола.'),
    ('4 соседа: (1,0)(-1,0)(0,1)(0,-1)', 'Только вверх/вниз/влево/вправо, не по диагонали.'),
    ('if достигли goal — break', 'Путь найден.'),
    ('if prev[goal]==-1 return from', 'Игрок в недостижимой зоне.'),
    ('идём от goal назад до start', 'Восстанавливаем путь.'),
    ('return первая клетка после start', 'Именно ОДИН шаг BFS — не весь путь сразу.'),
]

LEVELDATA_EXAMPLE = [
    ('M{ 12, 5, 1, 0, Chase, 1.05f, true }',
     'Клетка (12,5), патруль вправо, chase, скорость ×1.05, lockPatrolAxis.'),
    ('M{ 7, 13, 1, 0, Chase, 1.05f, true }',
     'Ур.3: злой на зелёной горизонтальной линии.'),
    ('M{ 6, 14, 0, -1, Patrol, 0.9f }',
     'Ур.1: патруль вверх по коридору, без lock.'),
]


def main() -> None:
    doc = Document()
    setup(doc)

    t = doc.add_heading("NPC (розовые монстры) — максимально подробно", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(doc,
      "Документ для тех, кто не понял с первого раза. Сначала — общая картина, "
      "потом разбор каждой строки кода. Файлы: Monster.hpp, Monster.cpp, "
      "MonsterPinkSheet, LevelData.hpp, фрагменты Game.cpp и Maze.cpp.")

    # OVERVIEW
    h1(doc, "Часть 0. Общая картина — что происходит от А до Я")
    logic_block(doc,
                "1) В LevelData.hpp для каждого уровня записан список монстров: где встать, "
                "куда патрулировать, Patrol или Chase. 2) Game::startLevel() создаёт объекты Monster "
                "и вызывает spawn(). 3) Каждый кадр (60–120 раз в секунду) Game вызывает monster.update() — "
                "монстр решает куда идти и меняет кадр анимации. 4) Game вызывает monster.draw() — "
                "на экран попадает кусок PNG из monster_pink_sheet.png. 5) Если круг игрока пересёкся с кругом "
                "монстра — startKill(), состояние Caught, потом смерть игрока.")
    p(doc,
      "У одного монстра в памяти есть: позиция (pos_), скорость, направление патруля, "
      "флаги chasing_/evil_, состояние анимации и номер кадра. Всё это — поля класса Monster в Monster.hpp.")

    h2(doc, "Схема состояний анимации (Chase-монстр)")
    p(doc, "Run (бег) → [игрок близко] → Transform (5 кадров) → Evil (злой, рот по дистанции) "
          "→ [игрок далеко] → снова Run. При касании игрока → Kill (4 кадра) → KillDone.")
    p(doc, "Patrol-монстр всегда в Run (кроме kill при столкновении).")

    h2(doc, "Схема ИИ Chase")
    p(doc, "Каждый кадр: измерить dist до игрока → обновить chasing_ (220/140) → "
          "если не chasing — updatePatrol(); если chasing — BFS один шаг к игроку.")

    # MONSTER.HPP
    h1(doc, "Часть 1. Monster.hpp — заголовок класса (каждая строка)")
    add_lines(doc, MONSTER_HPP)

    # PINK SHEET
    h1(doc, "Часть 2. MonsterPinkSheet — как устроен спрайт-лист")
    logic_block(doc,
                "Файл monster_pink_sheet.png — одна большая картинка 6×3 ячеек. "
                "cellRect(2, 1) вернёт прямоугольник 3-го кадра среднего ряда (превращение). "
                "Исходник monster_spritesheet.png нарезает scripts/extract_monster_pink_sheet.py.")
    add_lines(doc, PINK_SHEET, "MonsterPinkSheet.hpp / .cpp")

    # MONSTER.CPP
    h1(doc, "Часть 3. Monster.cpp — весь код по функциям")

    h2(doc, "3.1. Начало файла — константы и dirFromPatrol")
    add_lines(doc, MONSTER_CPP_HEADER)

    h2(doc, "3.2. void Monster::spawn(...) — рождение монстра на уровне")
    logic_block(doc,
                "Вызывается один раз при старте уровня из Game.cpp. "
                "cfg — одна запись M{...} из LevelData. baseSpeed — скорость патруля с уровня.")
    add_lines(doc, SPAWN_LINES)

    h2(doc, "3.3. startKill, onChaseStarted, onChaseCalmed, mouthColFromDist")
    add_lines(doc, [
        ('startKill: if уже Kill → return', 'Не перезапускать анимацию.'),
        ('animState_ = Kill; frame=0', 'Начать ряд ROW_KILL с первого кадра.'),
        ('onChaseStarted: только если Chase и лист загружен', 'Иначе выход.'),
        ('animState_ = Transform', 'Средний ряд — «злится».'),
        ('onChaseCalmed: evil_=false; state=Run', 'Успокоился — снова обычный бег.'),
        ('mouthColFromDist: dist>220→0, >170→1, >130→2, >90→3, иначе 4',
         'Чем ближе игрок — тем «шире рот» (больше col в ROW_TRANSFORM).'),
    ])

    h2(doc, "3.4. void Monster::updatePatrol — как ходит патруль")
    logic_block(doc, "Используется и у Patrol-монстров, и у Chase когда игрок далеко.")
    add_lines(doc, PATROL_LINES)

    h2(doc, "3.5. void Monster::updateChase — погоня за игроком")
    logic_block(doc, "Самая сложная часть ИИ. Ключ — nextStepBfs: монстр не летит сквозь стены, "
                      "а идёт по коридорам как в tower defense.")
    add_lines(doc, CHASE_LINES)

    h2(doc, "3.6. void Monster::update — главный метод кадра")
    add_lines(doc, UPDATE_LINES)

    h2(doc, "3.7. void Monster::advanceAnim — смена кадров PNG")
    add_lines(doc, ADVANCE_ANIM)

    h2(doc, "3.8. void Monster::draw и drawSprite — как рисуем на экране")
    add_lines(doc, DRAW_LINES)

    # GAME
    h1(doc, "Часть 4. Как Game.cpp подключает монстров")
    add_lines(doc, GAME_NPC)

    h2(doc, "Фрагмент Game::update (столкновение) — построчно")
    add_lines(doc, [
        ('for (int i = 0; i < monsters_.size(); ++i)', 'Проверяем каждого монстра.'),
        ('const auto& monster = monsters_[i]', 'Ссылка на i-го монстра.'),
        ('circlesOverlap(player pos/radius, monster pos/radius)',
         'dx²+dy² < (r1+r2)² — круги пересеклись.'),
        ('playDefeatSound()', 'Останавливаем музыку, звук поражения.'),
        ('if (pinkSheet loaded)', 'Если есть спрайты — полная анимация kill.'),
        ('catchingMonster_ = i', 'Запоминаем индекс «съевшего» монстра.'),
        ('monsters_[i].startKill()', 'Запуск kill-анимации.'),
        ('state_ = GameState::Caught', 'Игрок скрыт, виден только монстр.'),
        ('else → сразу Dying', 'Без спрайтов — сразу смерть игрока.'),
    ])

    h2(doc, "Состояние Caught в Game::update")
    add_lines(doc, [
        ('monsters_[catchingMonster_].updateKill(dt)', 'Крутим только kill-анимацию.'),
        ('if (killFinished())', '4 кадра kill проиграны.'),
        ('player_.startDeath()', 'Начинаем анимацию смерти игрока.'),
        ('state_ = Dying', 'Переход к смерти.'),
    ])

    # MAZE BFS
    h1(doc, "Часть 5. Maze::nextStepBfs — как монстр «знает» путь")
    logic_block(doc,
                "Лабиринт — сетка. BFS заливает волну от монстра до клетки игрока только по '.' клеткам. "
                "Потом от клетки игрока идёт назад к старту и берётся первая клетка на пути — "
                "это и есть следующий шаг. Каждый кадр — новый BFS (игрок движется — путь пересчитывается).")
    add_lines(doc, BFS_LINES)

    h2(doc, "Maze::moveCircle — почему монстр не застревает в углах")
    add_lines(doc, [
        ('cr = radius * 0.88', 'Чуть меньший радиус коллизии — плавнее углы.'),
        ('пробуем newX = pos.x + delta.x', 'Сначала только горизонталь.'),
        ('if (!isWallAtPixel(newX, pos.y, cr)) pos.x = newX', 'Если не в стене — принимаем X.'),
        ('потом то же для Y', 'В углу можно проскользнуть вдоль стены.'),
        ('isWallAtPixel — круг vs прямоугольник клетки #', 'Проверка по всем соседним клеткам сетки.'),
    ])

    # LEVEL DATA
    h1(doc, "Часть 6. LevelData.hpp — как задаём монстров на уровне")
    logic_block(doc,
                "buildLevels() возвращает vector<LevelConfig>. У каждого уровня поле monsters — "
                "vector<MonsterSpawn>. Буква M — сокращение для MonsterSpawn{...}.")
    add_lines(doc, LEVELDATA_EXAMPLE, "Примеры записей M{...}")

    h2(doc, "Расшифровка полей MonsterSpawn")
    add_lines(doc, [
        ('gridX, gridY', 'Координата клетки сетки (счёт с 0, слева-верх).'),
        ('patrolDirX, patrolDirY', 'Один из {-1,0,1} — направление патруля.'),
        ('MonsterAI::Patrol / Chase', 'Тип ИИ.'),
        ('speedMultiplier', '1.0 = норма, 1.15 = на 15% быстрее.'),
        ('lockPatrolAxis (опционально)', 'true = не сворачивать на перекрёстках.'),
    ])

    # EXAMPLE WALKTHROUGH
    h1(doc, "Часть 7. Пример: один кадр жизни Chase-монстра")
    p(doc, "Допустим, ур.5, монстр M{19,11,1,0,Chase,...}, игрок подбежал на 150 px.")
    bullet = doc.add_paragraph
    steps = [
        "Game::update вызывает monster.update(0.016, maze, playerPos, 130).",
        "updateChase: dist=150 < 220 → chasing_=true.",
        "wasChasing был false → onChaseStarted() → Transform, frame 0.",
        "Пока Transform — return без движения несколько кадров.",
        "После 5 кадров Transform → evil_=true, Evil.",
        "pixelToGrid: монстр (19,11), игрок (25,11).",
        "nextStepBfs → клетка (20,11). gridCenter → пиксели центра.",
        "dir = вправо, moveCircle сдвигает pos_.",
        "lastMoving_=true, advanceAnim в Evil крутит animFrame_.",
        "drawSprite: row=TRANSFORM, col=mouthColFromDist(150)=2 (рот средне открыт).",
        "setScale(-scale или +scale) по facingLeft_.",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}", style="List Number")

    h1(doc, "Часть 8. Откуда взялась картинка монстра")
    p(doc, "1) Исходный PNG monster_spritesheet.png в assets/. "
          "2) Python находит прозрачные промежутки между кадрами. "
          "3) Собирает ровную сетку 6×3 → monster_pink_sheet.png. "
          "4) Art.cpp загружает в art_.monsterPinkSheet. "
          "5) Каждому Monster передаётся указатель setPinkSheet(&art_.monsterPinkSheet) — "
          "все делят одну текстуру, но у каждого свои pos_, animFrame_, animState_.")
    p(doc, "Пересборка листа: py -3 scripts/extract_monster_pink_sheet.py")

    h1(doc, "Часть 9. Шпаргалка: какой файл за что")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Вопрос"
    tbl.rows[0].cells[1].text = "Где смотреть"
    qa = [
        ("Где задать позицию монстра?", "LevelData.hpp → M{ gridX, gridY, ... }"),
        ("Как включить погоню?", "MonsterAI::Chase в M{...}"),
        ("Почему монстр сходит с линии?", "Добавить lockPatrolAxis: true"),
        ("Как меняется скорость погони?", "Constants.hpp MONSTER_CHASE_FACTOR, updateChase"),
        ("Как меняется рот?", "mouthColFromDist + drawSprite Evil"),
        ("Как монстр обходит стены?", "Maze::nextStepBfs + moveCircle"),
        ("Что при касании игрока?", "Game.cpp circlesOverlap → startKill → Caught"),
        ("Сколько кадров бега?", "MonsterPinkSheet::RUN_FRAMES = 6"),
    ]
    for q, a in qa:
        r = tbl.add_row().cells
        r[0].text = q
        r[1].text = a

    # FULL Monster.cpp line by line
    h1(doc, "Часть 10. Monster.cpp — ПОЛНЫЙ построчный разбор (все 280 строк)")
    p(doc,
      "Ниже — каждая строка файла src/Monster.cpp с пояснением. "
      "Пустые строки и закрывающие скобки тоже указаны для навигации по файлу.")

    cpp_path = ROOT / "src" / "Monster.cpp"
    explanations = _monster_cpp_explanations()
    lines = cpp_path.read_text(encoding="utf-8").splitlines()
    full_rows: list[tuple[str, str]] = []
    for lineno, line in enumerate(lines, 1):
        stripped = line.strip()
        key = stripped
        expl = explanations.get(key) or explanations.get(f"L{lineno}") or _guess_expl(stripped, lineno)
        display = line if len(line) <= 90 else line[:87] + "..."
        full_rows.append((f"{lineno}: {display}", expl))

    # split into chunks for readability (word tables get huge)
    chunk = 35
    for i in range(0, len(full_rows), chunk):
        h3(doc, f"Строки {i + 1}–{min(i + chunk, len(full_rows))} файла Monster.cpp")
        add_lines(doc, full_rows[i : i + chunk])

    p(doc, "— Конец документа. Откройте src/Monster.cpp рядом с этим файлом для сверки. —")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


def _guess_expl(stripped: str, lineno: int) -> str:
    if not stripped:
        return "Пустая строка — разделение блоков кода для читаемости."
    if stripped in ("{", "}", "} // namespace"):
        return "Открытие/закрытие блока { }."
    if stripped.startswith("//"):
        return "Комментарий в коде (в Monster.cpp комментариев мало)."
    if stripped == "return;":
        return "Выход из функции — дальнейший код этого кадра не выполняется."
    if stripped == "return":
        return "Выход из функции (void)."
    return "См. пояснения к функции выше в Части 3 или контекст соседних строк."


def _monster_cpp_explanations() -> dict[str, str]:
    """Точные пояснения для ключевых строк Monster.cpp."""
    return {
        '#include "Monster.hpp"': "Подключаем объявление класса Monster.",
        "#include <algorithm>": "Библиотека std::min — ограничиваем индекс кадра.",
        "#include <cmath>": "Математика: sqrt, abs.",
        '#include "Constants.hpp"': "Глобальные константы CHASE_START_DISTANCE и др.",
        "constexpr float RUN_FRAME_TIME = 0.09f;": "0.09 сек ≈ 11 кадров бега в секунду.",
        "constexpr float TRANSFORM_FRAME_TIME = 0.14f;": "Медленнее — превращение заметнее.",
        "constexpr float KILL_FRAME_TIMES[] = {0.18f, 0.18f, 0.2f, 0.22f, 0.4f};":
            "Массив длительностей каждого из 5 таймингов kill (используем 4 кадра).",
        "sf::Vector2f dirFromPatrol(int px, int py) {": "Вспомогательная: int-направление → вектор.",
        "sf::Vector2f d(static_cast<float>(px), static_cast<float>(py));":
            "Вектор из целочисленного направления патруля.",
        "const float len = std::sqrt(d.x * d.x + d.y * d.y);": "Длина вектора до нормализации.",
        "if (len > 0.f) d /= len;": "Делим на длину — получаем единичный вектор (скорость не зависит от диагонали).",
        "return d;": "Возвращаем направление длиной 1.",
        "} // namespace": "Конец анонимного namespace — константы видны только в этом .cpp.",
        "void Monster::spawn(const MonsterSpawn& cfg, float baseSpeed, const Maze& maze) {":
            "Метод класса Monster — инициализация при появлении на уровне.",
        "pos_ = maze.nearestFloorCenter(cfg.gridX, cfg.gridY);":
            "Пиксельная позиция центра клетки; если стена — ищется ближайший пол.",
        "patrolDirX_ = cfg.patrolDirX;": "Копируем горизонтальное направление из LevelData.",
        "patrolDirY_ = cfg.patrolDirY;": "Копируем вертикальное направление.",
        "lockPatrolAxis_ = cfg.lockPatrolAxis;": "Флаг «только разворот на линии».",
        "ai_ = cfg.ai;": "Patrol или Chase.",
        "chasing_ = false;": "Ещё не в режиме погони.",
        "evil_ = false;": "Ещё не превратился в злого.",
        "facingLeft_ = patrolDirX_ < 0;": "Начальная ориентация спрайта.",
        "speedMultiplier_ = cfg.speedMultiplier;": "Индивидуальный множитель скорости монстра.",
        "animState_ = MonsterAnimState::Run;": "Стартовая анимация — обычный бег.",
        "animFrame_ = 0;": "Первый кадр ряда.",
        "animTimer_ = 0.f;": "Таймер кадра сброшен.",
        "lastMoving_ = false;": "Пока не двигался.",
        "playerDist_ = 9999.f;": "Игрок условно очень далеко.",
        "if (ai_ == MonsterAI::Patrol) {": "Ветка для патрульного типа.",
        "speed_ = baseSpeed * speedMultiplier_;": "Скорость с уровня × множитель.",
        "} else {": "Иначе — Chase.",
        "speed_ = 45.f * speedMultiplier_;": "Базовая скорость chase в режиме «спокойствия».",
        "void Monster::startKill() {": "Вызывается из Game при столкновении с игроком.",
        "if (animState_ == MonsterAnimState::Kill || animState_ == MonsterAnimState::KillDone)":
            "Уже убивает — не сбрасывать анимацию повторно.",
        "animState_ = MonsterAnimState::Kill;": "Переключаемся на ряд ROW_KILL.",
        "void Monster::onChaseStarted() {": "Момент, когда chasing_ стал true впервые.",
        "if (ai_ != MonsterAI::Chase || !pinkSheet_ || !pinkSheet_->loaded || evil_ ||":
            "Только chase, с текстурой, ещё не злой, не в kill/transform.",
        "animState_ == MonsterAnimState::Transform || animState_ == MonsterAnimState::Kill ||":
            "Не прерывать уже идущие особые анимации.",
        "animState_ == MonsterAnimState::KillDone)": "Конец условия guard.",
        "animState_ = MonsterAnimState::Transform;": "Запуск 5 кадров злости.",
        "void Monster::onChaseCalmed() {": "Игрок убежал — успокоение.",
        "if (!evil_ && animState_ != MonsterAnimState::Evil &&":
            "Если и так не злой — делать нечего.",
        "animState_ != MonsterAnimState::Transform)": "И не в процессе transform.",
        "evil_ = false;": "Сброс флага злости.",
        "animState_ = MonsterAnimState::Run;": "Вернуться к обычному бегу.",
        "int Monster::mouthColFromDist(float dist) const {": "Номер колонки рта по дистанции.",
        "if (dist > CHASE_START_DISTANCE) return 0;": "Дальше 220 — рот «закрыт» (col 0 не рисуем в evil).",
        "if (dist > 170.f) return 1;": "Далековато — слегка открыт.",
        "if (dist > 130.f) return 2;": "Средняя дистанция.",
        "if (dist > 90.f) return 3;": "Близко.",
        "return 4;": "Очень близко — максимально открытый рот.",
        "void Monster::updatePatrol(float dt, const Maze& maze) {": "Логика хождения туда-сюда.",
        "const sf::Vector2f wish = dirFromPatrol(patrolDirX_, patrolDirY_) * speed_ * dt;":
            "Вектор смещения за этот кадр: направление × скорость × время.",
        "const sf::Vector2f before = pos_;": "Запомнить позицию до хода.",
        "maze.moveCircle(pos_, wish, radius_);": "Сдвинуть с коллизией; pos_ изменится если проход свободен.",
        "if ((pos_ - before).x * (pos_ - before).x + (pos_ - before).y * (pos_ - before).y < 0.01f) {":
            "Квадрат смещения < 0.01 — практически не двинулись → уперлись в стену.",
        "if (lockPatrolAxis_) {": "Режим «не сходить с оси коридора».",
        "patrolDirX_ = -patrolDirX_;": "Разворот по X.",
        "patrolDirY_ = -patrolDirY_;": "Разворот по Y.",
        "const int options[4][2] = {{1, 0}, {-1, 0}, {0, 1}, {0, -1}};":
            "Четыре направления для поиска обхода на перекрёстке.",
        "for (const auto& opt : options) {": "Перебираем вправо, влево, вниз, вверх.",
        "sf::Vector2f testPos = pos_;": "Копия позиции для пробного шага.",
        "const sf::Vector2f tryWish = dirFromPatrol(opt[0], opt[1]) * speed_ * dt;":
            "Пробный вектор в одном из 4 направлений.",
        "maze.moveCircle(testPos, tryWish, radius_);": "Пробуем сдвинуть копию — стену не ломаем.",
        "if ((testPos - pos_).x * (testPos - pos_).x + (testPos - pos_).y * (testPos - pos_).y > 0.01f) {":
            "Если копия сдвинулась — направление проходимо.",
        "patrolDirX_ = opt[0];": "Принимаем новое направление X.",
        "patrolDirY_ = opt[1];": "Принимаем новое направление Y.",
        "pos_ = testPos;": "Применяем удачный пробный шаг.",
        "patrolDirX_ = -patrolDirX_;": "Крайний случай: просто разворот.",
        "patrolDirY_ = -patrolDirY_;": "Разворот Y.",
        "void Monster::updateChase(float dt, const Maze& maze, sf::Vector2f playerPos, float playerSpeed) {":
            "ИИ преследования — вызывается только для Chase из update().",
        "speed_ = playerSpeed * MONSTER_CHASE_FACTOR * speedMultiplier_;":
            "Скорость = скорость игрока × 0.82 × множитель — монстр чуть медленнее.",
        "const float dx = playerPos.x - pos_.x;": "Разница по X до игрока.",
        "const float dy = playerPos.y - pos_.y;": "Разница по Y.",
        "const float dist = std::sqrt(dx * dx + dy * dy);": "Прямое расстояние в пикселях (не по сетке).",
        "playerDist_ = dist;": "Сохраняем для drawSprite / рот.",
        "const bool wasChasing = chasing_;": "Запоминаем старый флаг погони.",
        "if (!chasing_ && dist < CHASE_START_DISTANCE) chasing_ = true;":
            "Включить погоню если игрок вошёл в радиус 220 px.",
        "if (chasing_ && dist > CHASE_STOP_DISTANCE) chasing_ = false;":
            "Выключить если отбежал дальше 140 px — иначе монстр мгновенно бросал бы погоню.",
        "if (!wasChasing && chasing_) onChaseStarted();":
            "Переход false→true — первый момент обнаружения игрока.",
        "if (!chasing_) {": "Игрок далеко — не преследуем.",
        "onChaseCalmed();": "Если был злой — успокоиться.",
        "updatePatrol(dt, maze);": "Ходим как патруль.",
        "if (animState_ == MonsterAnimState::Transform)": "Во время превращения.",
        "const sf::Vector2i fromG = maze.pixelToGrid(pos_);": "Клетка монстра.",
        "const sf::Vector2i toG = maze.pixelToGrid(playerPos);": "Клетка игрока.",
        "const sf::Vector2i nextG = maze.nextStepBfs(fromG, toG);": "Следующий шаг BFS по полу.",
        "if (nextG == fromG) return;": "Некуда идти.",
        "const sf::Vector2f target = maze.gridCenter(nextG);": "Центр следующей клетки в px.",
        "sf::Vector2f dir(target.x - pos_.x, target.y - pos_.y);": "Вектор к цели.",
        "const float len = std::sqrt(dir.x * dir.x + dir.y * dir.y);": "Длина вектора.",
        "if (len > 0.001f) dir /= len;": "Нормализация — только направление.",
        "maze.moveCircle(pos_, dir * speed_ * dt, radius_);": "Шаг к игроку.",
        "void Monster::update(float dt, const Maze& maze, sf::Vector2f playerPos, float playerSpeed) {":
            "Вызывается Game каждый кадр для каждого монстра.",
        "if (animState_ == MonsterAnimState::Kill || animState_ == MonsterAnimState::KillDone) {":
            "Во время kill не ходим.",
        "advanceAnim(dt, false);": "Только кадры kill; moving=false.",
        "const sf::Vector2f before = pos_;": "Для lastMoving_.",
        "if (ai_ == MonsterAI::Patrol) {": "Патрульный тип.",
        "updatePatrol(dt, maze);": "Простое хождение.",
        "const float pdx = playerPos.x - pos_.x;": "Дистанция до игрока (для отладки/рот не нужен).",
        "playerDist_ = std::sqrt(pdx * pdx + pdy * pdy);": "Обновить dist.",
        "updateChase(dt, maze, playerPos, playerSpeed);": "Chase ИИ.",
        "const sf::Vector2f delta = pos_ - before;": "Насколько сдвинулись за кадр.",
        "const float movedSq = delta.x * delta.x + delta.y * delta.y;": "Квадрат смещения.",
        "lastMoving_ = movedSq > 0.25f;": ">0.5 px — считаем движением.",
        "if (lastMoving_ && std::abs(delta.x) > 0.05f)": "Если двигались по горизонтали.",
        "facingLeft_ = delta.x < 0.f;": "Смотрим в сторону движения.",
        "else if (patrolDirX_ != 0)": "Иначе ориентация по патрулю.",
        "facingLeft_ = patrolDirX_ < 0;": "Патруль влево.",
        "advanceAnim(dt, lastMoving_);": "Смена кадра спрайта.",
        "void Monster::updateKill(float dt) {": "Отдельный вызов из Game в состоянии Caught.",
        "if (animState_ != MonsterAnimState::Kill) return;": "Только если в kill.",
        "void Monster::advanceAnim(float dt, bool moving) {": "Таймеры кадров для всех animState.",
        "if (!pinkSheet_ || !pinkSheet_->loaded) return;": "Без PNG нечего анимировать.",
        "animTimer_ += dt;": "Прибавить время кадра.",
        "if (animState_ == MonsterAnimState::Run) {": "Обычный бег.",
        "if (!moving) {": "Стоит на месте.",
        "animFrame_ = 0;": "Первый кадр = поза покоя.",
        "animTimer_ = 0.f;": "Сброс таймера.",
        "if (animTimer_ >= RUN_FRAME_TIME) {": "Пора сменить кадр бега.",
        "animTimer_ -= RUN_FRAME_TIME;": "Остаток времени в следующий кадр.",
        "animFrame_ = (animFrame_ + 1) % MonsterPinkSheet::RUN_FRAMES;": "0→1→2→3→4→5→0…",
        "if (animState_ == MonsterAnimState::Evil) {": "Злой режим.",
        "if (moving && animTimer_ >= RUN_FRAME_TIME) {": "При движении — мелькают кадры бега под ртом.",
        "if (animState_ == MonsterAnimState::Transform) {": "Превращение.",
        "if (animTimer_ >= TRANSFORM_FRAME_TIME) {": "Пора следующий кадр transform.",
        "animTimer_ -= TRANSFORM_FRAME_TIME;": "Сброс с остатком.",
        "++animFrame_;": "Следующий кадр transform.",
        "if (animFrame_ >= MonsterPinkSheet::TRANSFORM_FRAMES) {": "Все 5 кадров прошли.",
        "animFrame_ = MonsterPinkSheet::TRANSFORM_FRAMES - 1;": "Зафиксировать последний кадр transform.",
        "evil_ = true;": "Теперь официально злой.",
        "animState_ = MonsterAnimState::Evil;": "Переключить режим анимации.",
        "if (animState_ == MonsterAnimState::Kill) {": "Анимация поимки.",
        "const int frameIdx = std::min(animFrame_, MonsterPinkSheet::KILL_PLAY_FRAMES - 1);":
            "Индекс не выходит за 0..3.",
        "if (animTimer_ >= KILL_FRAME_TIMES[frameIdx]) {": "Дождаться времени этого кадра.",
        "animTimer_ -= KILL_FRAME_TIMES[frameIdx];": "Остаток времени.",
        "++animFrame_;": "Следующий кадр kill.",
        "if (animFrame_ >= MonsterPinkSheet::KILL_PLAY_FRAMES) {": "Все 4 кадра показаны.",
        "animFrame_ = MonsterPinkSheet::KILL_PLAY_FRAMES - 1;": "Держим последний кадр.",
        "animState_ = MonsterAnimState::KillDone;": "Game узнает что killFinished().",
        "void Monster::draw(sf::RenderTarget& target) const {": "Рисуем каждый кадр в render().",
        "if (pinkSheet_ && pinkSheet_->loaded) {": "Если PNG загружен.",
        "drawSprite(target);": "Спрайт с листа.",
        "sf::CircleShape body(radius_);": "Заглушка — круг.",
        "body.setOrigin(radius_, radius_);": "Центр круга в pos_.",
        "body.setPosition(pos_);": "Позиция на карте.",
        "body.setFillColor((chasing_ || evil_) ? sf::Color(240, 90, 90) : sf::Color(220, 120, 60));":
            "Красный если злой/погоня, иначе оранжевый.",
        "body.setOutlineThickness(2.f);": "Обводка.",
        "body.setOutlineColor(sf::Color(40, 20, 20));": "Тёмная обводка.",
        "target.draw(body);": "Отправить в окно.",
        "void Monster::drawSprite(sf::RenderTarget& target) const {": "Выбор кадра и отрисовка PNG.",
        "int row = MonsterPinkSheet::ROW_RUN;": "По умолчанию верхний ряд.",
        "int col = 0;": "По умолчанию первый кадр.",
        "if (animState_ == MonsterAnimState::Run) {": "Обычный бег.",
        "row = MonsterPinkSheet::ROW_RUN;": "Верхний ряд.",
        "col = animFrame_;": "Текущий кадр бега 0..5.",
        "} else if (animState_ == MonsterAnimState::Transform) {": "Превращение.",
        "row = MonsterPinkSheet::ROW_TRANSFORM;": "Средний ряд.",
        "col = animFrame_;": "Кадр 0..4 transform.",
        "} else if (animState_ == MonsterAnimState::Evil) {": "Злой.",
        "col = mouthColFromDist(playerDist_);": "Колонка по дистанции.",
        "if (col < 1) col = 1;": "Минимум col=1 — не показывать «закрытый» рот в evil.",
        "if (lastMoving_ && animFrame_ % 2 == 1 && col < MonsterPinkSheet::TRANSFORM_FRAMES - 1)":
            "При беге на нечётном кадре — чуть шире рот (живость).",
        "++col;": "Сдвиг колонки рта.",
        "} else if (animState_ == MonsterAnimState::Kill || animState_ == MonsterAnimState::KillDone) {":
            "Поимка.",
        "row = MonsterPinkSheet::ROW_KILL;": "Нижний ряд.",
        "col = std::min(animFrame_, MonsterPinkSheet::KILL_PLAY_FRAMES - 1);": "Кадр kill 0..3.",
        "sf::Sprite sprite(pinkSheet_->texture);": "SFML спрайт с общей текстурой.",
        "sprite.setTextureRect(pinkSheet_->cellRect(col, row));": "Вырезать нужную ячейку.",
        "const float cw = static_cast<float>(pinkSheet_->cellW);": "Ширина ячейки.",
        "const float ch = static_cast<float>(pinkSheet_->cellH);": "Высота ячейки.",
        "const float displayH = radius_ * 4.0f;": "Высота на экране = 4× радиус (48 px при r=12).",
        "const float scale = displayH / ch;": "Масштаб от высоты ячейки.",
        "sprite.setOrigin(cw * 0.5f, ch * 0.55f);": "Опорная точка — центр по X, 55% по Y (ноги).",
        "sprite.setScale(facingLeft_ ? -scale : scale, scale);":
            "Отрицательный X = отзеркалить влево; Y всегда положительный.",
        "sprite.setPosition(pos_);": "Координаты на карте.",
        "target.draw(sprite);": "Нарисовать в окно.",
    }


if __name__ == "__main__":
    main()
